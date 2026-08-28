import tempfile
import unittest
from zipfile import ZipFile

from docx import Document

from backend.docx_exporter import export_html_to_docx
from backend.renderer import render_markdown


def _document_text(document: Document) -> str:
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(texts)


class TestDocxExporter(unittest.TestCase):
    def test_export_skips_whitespace_between_block_elements(self):
        html = """<body>
<p>first</p>
<h2>heading</h2>
<ul><li>item</li></ul>
<pre><code>code</code></pre>
<p>last</p>
</body>"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = f"{tmp_dir}/export.docx"

            export_html_to_docx(html, output_path)

            document = Document(output_path)
            self.assertEqual(
                [paragraph.text for paragraph in document.paragraphs],
                ["first", "heading", "item", "last"],
            )
            self.assertIn("code", _document_text(document))

    def test_export_preserves_clickable_links(self):
        html = render_markdown(
            "[Project documentation](https://example.com/docs)\n\n"
            "Bare URL: https://example.com/bare."
        )
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = f"{tmp_dir}/links.docx"

            export_html_to_docx(html, output_path)

            document = Document(output_path)
            body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("Project documentation", body_text)
            self.assertIn("https://example.com/bare", body_text)

            with ZipFile(output_path) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
                relationships_xml = archive.read("word/_rels/document.xml.rels").decode(
                    "utf-8"
                )

            self.assertEqual(document_xml.count("<w:hyperlink"), 2)
            self.assertIn('Target="https://example.com/docs"', relationships_xml)
            self.assertIn('Target="https://example.com/bare"', relationships_xml)
            self.assertIn(
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"',
                relationships_xml,
            )

    def test_export_styles_code_blocks_with_background(self):
        html = render_markdown("```python\nprint('hello')\n```")
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = f"{tmp_dir}/code.docx"

            export_html_to_docx(html, output_path)

            document = Document(output_path)
            self.assertIn("print", _document_text(document))

            with ZipFile(output_path) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")

            self.assertIn('<w:shd w:fill="F6F8FA"/>', document_xml)
            self.assertIn("<w:tcMar>", document_xml)
            self.assertIn('<w:left w:w="120" w:type="dxa"/>', document_xml)
            self.assertIn('<w:right w:w="120" w:type="dxa"/>', document_xml)
            self.assertNotIn("<w:ind ", document_xml)
            self.assertNotIn('w:before="120"', document_xml)
            self.assertNotIn('w:after="120"', document_xml)
            self.assertIn('w:ascii="Courier New"', document_xml)

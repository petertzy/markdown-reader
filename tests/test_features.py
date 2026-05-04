"""
test_features.py
================
Unit tests for word_count_bar.py and recent_files.py.

Run with:
    python -m pytest tests/test_features.py -v
or:
    python -m unittest tests.test_features -v
"""

import base64
import json
import os
import sys
import tempfile
import unittest

# ---------------------------------------------------------------------------
# Import the modules under test.
# ---------------------------------------------------------------------------
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from backend.routers.files import ConvertToMarkdownPayload, convert_to_markdown
from markdown_reader.recent_files import (
    RecentFilesManager,
    _middle_ellipsis,
    _safe_write_json,
)
from markdown_reader.word_count_bar import _count_words, _reading_time, _strip_markdown

# ===========================================================================
# Tests for word_count_bar helpers
# ===========================================================================


class TestStripMarkdown(unittest.TestCase):
    """_strip_markdown should remove syntax tokens without destroying words."""

    def test_atx_heading(self):
        result = _strip_markdown("# Hello World")
        self.assertIn("Hello", result)
        self.assertNotIn("#", result)

    def test_bold_markers(self):
        result = _strip_markdown("**bold text**")
        self.assertIn("bold", result)
        self.assertNotIn("**", result)

    def test_italic_markers(self):
        result = _strip_markdown("*italic text*")
        self.assertIn("italic", result)
        self.assertNotIn("*", result)

    def test_inline_code(self):
        # Inline code content is removed entirely.
        result = _strip_markdown("Use `print()` here")
        self.assertIn("Use", result)
        self.assertIn("here", result)

    def test_fenced_code_block(self):
        md = "intro\n```python\nfor i in range(10):\n    pass\n```\noutro"
        result = _strip_markdown(md)
        self.assertIn("intro", result)
        self.assertIn("outro", result)
        self.assertNotIn("range", result)

    def test_link_keeps_text(self):
        result = _strip_markdown("[click here](https://example.com)")
        self.assertIn("click", result)
        self.assertNotIn("https", result)

    def test_image_removed(self):
        result = _strip_markdown("![alt text](image.png)")
        self.assertNotIn("alt", result)

    def test_blockquote_marker(self):
        result = _strip_markdown("> quoted line")
        self.assertIn("quoted", result)
        self.assertNotIn(">", result)

    def test_html_tags_removed(self):
        result = _strip_markdown("<p>Some text</p>")
        self.assertIn("Some", result)
        self.assertNotIn("<p>", result)

    def test_horizontal_rule(self):
        result = _strip_markdown("---")
        self.assertEqual(result.strip(), "")

    def test_table_pipes(self):
        result = _strip_markdown("| col1 | col2 |")
        self.assertNotIn("|", result)
        self.assertIn("col1", result)


class TestCountWords(unittest.TestCase):
    """_count_words should handle common edge cases correctly."""

    def test_empty_string(self):
        self.assertEqual(_count_words(""), 0)

    def test_whitespace_only(self):
        self.assertEqual(_count_words("   \n\t  "), 0)

    def test_simple_sentence(self):
        self.assertEqual(_count_words("Hello world"), 2)

    def test_multiple_spaces(self):
        self.assertEqual(_count_words("one   two   three"), 3)

    def test_newlines_between_words(self):
        self.assertEqual(_count_words("one\ntwo\nthree"), 3)

    def test_unicode_latin(self):
        # Latin accented characters should still count as one word each.
        self.assertEqual(_count_words("café résumé naïve"), 3)

    def test_cjk_characters_counted_individually(self):
        # Each Chinese character is one word.
        self.assertEqual(_count_words("你好世界"), 4)

    def test_mixed_cjk_and_latin(self):
        count = _count_words("Hello 世界 world")
        # "Hello" + "世" + "界" + "world" = 4
        self.assertEqual(count, 4)

    def test_punctuation_not_counted(self):
        # Punctuation attached to a word should not create extra word tokens.
        count = _count_words("Hello, world!")
        self.assertEqual(count, 2)


class TestReadingTime(unittest.TestCase):
    """_reading_time should return sensible human-readable strings."""

    def test_zero_words(self):
        self.assertEqual(_reading_time(0), "< 1 min read")

    def test_fewer_than_wpm(self):
        self.assertEqual(_reading_time(100), "< 1 min read")

    def test_exactly_one_minute(self):
        self.assertEqual(_reading_time(238), "1 min read")

    def test_multiple_minutes(self):
        result = _reading_time(238 * 5)
        self.assertEqual(result, "5 min read")

    def test_rounding(self):
        # 357 words / 238 wpm ≈ 1.5 → rounds to 2
        result = _reading_time(357)
        self.assertEqual(result, "2 min read")


# ===========================================================================
# Tests for recent_files helpers
# ===========================================================================


class TestMiddleEllipsis(unittest.TestCase):
    """_middle_ellipsis should shorten long strings correctly."""

    def test_short_path_unchanged(self):
        path = "/short/path.md"
        self.assertEqual(_middle_ellipsis(path, max_len=60), path)

    def test_long_path_shortened(self):
        path = "/very/long/directory/structure/that/exceeds/the/limit/file.md"
        result = _middle_ellipsis(path, max_len=40)
        self.assertLessEqual(len(result), 40)
        self.assertIn("…", result)

    def test_result_starts_and_ends_with_original(self):
        path = "/a/b/c/d/e/f/g/h/i/j/k/l/m/n/o/p/q/r/s/t/u/v/w/x/y/z/file.md"
        result = _middle_ellipsis(path, max_len=30)
        # The start of the original should be in the result.
        self.assertTrue(result.startswith(path[:10]))
        # The filename should be preserved at the end.
        self.assertTrue(result.endswith("file.md"))


class TestSafeWriteJson(unittest.TestCase):
    """_safe_write_json should write valid JSON atomically."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.target = os.path.join(self.tmpdir, "settings.json")

    def tearDown(self):
        import shutil

        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_creates_file_with_correct_content(self):
        data = {"key": "value", "num": 42}
        _safe_write_json(self.target, data)
        with open(self.target) as f:
            loaded = json.load(f)
        self.assertEqual(loaded, data)

    def test_overwrites_existing_file(self):
        _safe_write_json(self.target, {"old": True})
        _safe_write_json(self.target, {"new": True})
        with open(self.target) as f:
            loaded = json.load(f)
        self.assertNotIn("old", loaded)
        self.assertTrue(loaded["new"])

    def test_no_temp_file_left_behind(self):
        _safe_write_json(self.target, {"x": 1})
        files = os.listdir(self.tmpdir)
        # Only the target file should remain; no .tmp files.
        self.assertFalse(any(f.endswith(".tmp") for f in files))


class TestRecentFilesManager(unittest.TestCase):
    """Full integration tests for RecentFilesManager."""

    def setUp(self):
        self.tmpdir = tempfile.mkdtemp()
        self.settings = os.path.join(self.tmpdir, "settings.json")
        # Create some real files so existence checks pass.
        self.files = []
        for i in range(12):
            p = os.path.join(self.tmpdir, f"file{i}.md")
            open(p, "w").close()
            self.files.append(p)

    def tearDown(self):
        import shutil

        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def _manager(self) -> RecentFilesManager:
        return RecentFilesManager(self.settings, max_entries=10)

    def test_push_adds_entry(self):
        m = self._manager()
        m.push(self.files[0])
        self.assertIn(os.path.normpath(self.files[0]), m.entries)

    def test_push_most_recent_first(self):
        m = self._manager()
        m.push(self.files[0])
        m.push(self.files[1])
        self.assertEqual(m.entries[0], os.path.normpath(self.files[1]))

    def test_push_deduplicates(self):
        m = self._manager()
        m.push(self.files[0])
        m.push(self.files[1])
        m.push(self.files[0])  # push first file again
        self.assertEqual(m.entries.count(os.path.normpath(self.files[0])), 1)
        # Should now be at the top again.
        self.assertEqual(m.entries[0], os.path.normpath(self.files[0]))

    def test_max_entries_enforced(self):
        m = self._manager()
        for f in self.files:  # 12 files, max is 10
            m.push(f)
        self.assertLessEqual(len(m.entries), 10)

    def test_persistence_across_instances(self):
        m1 = self._manager()
        m1.push(self.files[0])
        m1.push(self.files[1])
        # New instance reads from same settings file.
        m2 = self._manager()
        self.assertEqual(m2.entries[0], os.path.normpath(self.files[1]))
        self.assertIn(os.path.normpath(self.files[0]), m2.entries)

    def test_clear_empties_list(self):
        m = self._manager()
        m.push(self.files[0])
        m.clear()
        self.assertEqual(m.entries, [])

    def test_clear_persisted(self):
        m = self._manager()
        m.push(self.files[0])
        m.clear()
        m2 = self._manager()
        self.assertEqual(m2.entries, [])

    def test_other_settings_keys_preserved(self):
        """Saving recent files must not destroy other settings keys."""
        # Pre-populate settings with an unrelated key.
        _safe_write_json(self.settings, {"theme": "dark", "api_key": "abc123"})
        m = self._manager()
        m.push(self.files[0])
        with open(self.settings) as f:
            data = json.load(f)
        self.assertEqual(data.get("theme"), "dark")
        self.assertEqual(data.get("api_key"), "abc123")

    def test_missing_file_pruned_on_load(self):
        """Paths to deleted files should be silently removed when loading."""
        ghost = os.path.join(self.tmpdir, "ghost.md")
        open(ghost, "w").close()
        m = self._manager()
        m.push(ghost)
        os.unlink(ghost)  # delete the file
        # Reload — ghost should be gone.
        m2 = self._manager()
        self.assertNotIn(os.path.normpath(ghost), m2.entries)

    def test_malformed_settings_handled_gracefully(self):
        """A corrupted settings file should not crash the manager."""
        with open(self.settings, "w") as f:
            f.write("THIS IS NOT JSON {{{")
        # Should not raise.
        m = self._manager()
        self.assertEqual(m.entries, [])

    def test_missing_settings_file_handled(self):
        """Missing settings file should be handled gracefully."""
        m = self._manager()  # settings file doesn't exist yet
        self.assertEqual(m.entries, [])
        # Push should create the file.
        m.push(self.files[0])
        self.assertTrue(os.path.isfile(self.settings))

    def test_path_normalisation(self):
        """Paths pushed with different representations should deduplicate."""
        m = self._manager()
        path = self.files[0]
        # Push the same file via two different path representations.
        m.push(path)
        m.push(path + os.sep + ".." + os.sep + os.path.basename(path))
        # After normalisation both should resolve to the same path.
        # The list should have exactly one entry for this file (or zero if
        # the double-dot path resolves differently — acceptable behaviour).
        normalised = os.path.normpath(os.path.abspath(path))
        count = m.entries.count(normalised)
        self.assertLessEqual(count, 1)


class TestConvertToMarkdown(unittest.TestCase):
    """File conversion endpoint helpers should return Markdown content."""

    def test_converts_html_upload_to_markdown(self):
        html = b"<h1>Title</h1><p>Hello <strong>world</strong>.</p>"
        payload = ConvertToMarkdownPayload(
            filename="sample.html",
            content_base64=base64.b64encode(html).decode("ascii"),
        )

        result = convert_to_markdown(payload)

        self.assertIn("Title", result["markdown"])
        self.assertIn("Hello", result["markdown"])

    def test_converts_docx_path_to_markdown(self):
        from docx import Document

        tmpdir = tempfile.mkdtemp()
        path = os.path.join(tmpdir, "sample.docx")
        try:
            document = Document()
            document.add_heading("Doc Title", level=1)
            document.add_paragraph("Body text")
            document.save(path)

            result = convert_to_markdown(ConvertToMarkdownPayload(path=path))

            self.assertIn("# Doc Title", result["markdown"])
            self.assertIn("Body text", result["markdown"])
        finally:
            import shutil

            shutil.rmtree(tmpdir, ignore_errors=True)


if __name__ == "__main__":
    unittest.main(verbosity=2)

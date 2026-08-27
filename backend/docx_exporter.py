from __future__ import annotations

import os
import re
from html import unescape
from html.parser import HTMLParser
from io import BytesIO
from typing import Any

import requests
from docx import Document
from docx.shared import Inches, Pt


def _resolve_image_source(src: str, base_dir: str | None) -> tuple[str, bytes | None]:
    src = src.strip()
    if src.startswith("file://"):
        src = src[7:]
    if re.match(r"^https?://", src, re.IGNORECASE):
        try:
            response = requests.get(src, timeout=15, allow_redirects=True)
            response.raise_for_status()
            return src, response.content
        except Exception:
            return src, None
    candidates = [src]
    if base_dir and not os.path.isabs(src):
        candidates.insert(0, os.path.abspath(os.path.join(base_dir, src)))
    for candidate in candidates:
        try:
            with open(candidate, "rb") as file_obj:
                return candidate, file_obj.read()
        except Exception:
            continue
    return src, None


class _DocxHtmlParser(HTMLParser):
    def __init__(self, document: Document, base_dir: str | None = None):
        super().__init__(convert_charrefs=False)
        self.document = document
        self.base_dir = base_dir
        self.current_paragraph = None
        self.current_table = None
        self.current_row = None
        self.current_cell = None
        self.current_list_style: list[str] = []
        self.current_style: dict[str, Any] = {
            "bold": False,
            "italic": False,
            "code": False,
        }
        self.data_buffer = ""

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str]]):
        attrs_dict = dict(attrs)
        self._flush_text()
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self.current_paragraph = self.document.add_heading("", level=int(tag[1]))
        elif tag in ("p", "blockquote", "pre"):
            self.current_paragraph = self.document.add_paragraph()
        elif tag == "code":
            self.current_style["code"] = True
        elif tag in ("strong", "b"):
            self.current_style["bold"] = True
        elif tag in ("em", "i"):
            self.current_style["italic"] = True
        elif tag in ("ul", "ol"):
            self.current_list_style.append(
                "List Bullet" if tag == "ul" else "List Number"
            )
        elif tag == "li":
            style = self.current_list_style[-1] if self.current_list_style else None
            self.current_paragraph = self.document.add_paragraph(style=style)
        elif tag == "br" and self.current_paragraph is not None:
            self.current_paragraph.add_run().add_break()
        elif tag == "img":
            src = attrs_dict.get("src", "")
            if not src:
                return
            image_path, image_data = _resolve_image_source(src, self.base_dir)
            if self.current_paragraph is None:
                self.current_paragraph = self.document.add_paragraph()
            try:
                if image_data is not None:
                    self.current_paragraph.add_run().add_picture(
                        BytesIO(image_data), width=Inches(6.0)
                    )
                else:
                    self.current_paragraph.add_run(f"[{image_path}]")
            except Exception:
                self.current_paragraph.add_run(f"[{image_path}]")
        elif tag == "table":
            self.current_table = []
        elif tag == "tr":
            self.current_row = []
        elif tag in ("td", "th"):
            self.current_cell = []

    def handle_endtag(self, tag: str):
        self._flush_text()
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6", "p", "blockquote", "li", "pre"):
            self.current_paragraph = None
        elif tag == "code":
            self.current_style["code"] = False
        elif tag in ("strong", "b"):
            self.current_style["bold"] = False
        elif tag in ("em", "i"):
            self.current_style["italic"] = False
        elif tag in ("ul", "ol") and self.current_list_style:
            self.current_list_style.pop()
        elif tag == "table":
            self._build_table()
            self.current_table = None
        elif (
            tag == "tr"
            and self.current_row is not None
            and self.current_table is not None
        ):
            self.current_table.append(self.current_row)
            self.current_row = None
        elif (
            tag in ("td", "th")
            and self.current_cell is not None
            and self.current_row is not None
        ):
            self.current_row.append("".join(self.current_cell).strip())
            self.current_cell = None

    def handle_data(self, data: str):
        self.data_buffer += data

    def _flush_text(self):
        if not self.data_buffer:
            return
        text = unescape(self.data_buffer)
        self.data_buffer = ""
        if self.current_cell is not None:
            self.current_cell.append(text)
            return
        if self.current_paragraph is None and not text.strip():
            return
        if self.current_paragraph is None:
            self.current_paragraph = self.document.add_paragraph()
        run = self.current_paragraph.add_run(text)
        run.bold = self.current_style.get("bold", False)
        run.italic = self.current_style.get("italic", False)
        if self.current_style.get("code", False):
            run.font.name = "Courier New"
            run.font.size = Pt(10)

    def _build_table(self):
        if not self.current_table:
            return
        rows = len(self.current_table)
        cols = max(len(row) for row in self.current_table)
        table = self.document.add_table(rows=rows, cols=cols)
        table.style = "Light Grid Accent 1"
        for row_index, row in enumerate(self.current_table):
            for col_index in range(cols):
                table.rows[row_index].cells[col_index].text = (
                    row[col_index] if col_index < len(row) else ""
                )


def export_html_to_docx(
    html_content: str, output_path: str, base_dir: str | None = None
) -> None:
    body_match = re.search(r"<body[^>]*>(.*)</body>", html_content, flags=re.S | re.I)
    body_html = body_match.group(1) if body_match else html_content
    document = Document()
    parser = _DocxHtmlParser(document, base_dir=base_dir)
    parser.feed(body_html)
    parser.close()
    document.save(output_path)

"""
backend/routers/export.py
=========================
File export endpoints (HTML, DOCX, PDF).
"""

from __future__ import annotations

import os
import sys
import tempfile

from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel

_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
if _ROOT not in sys.path:
    sys.path.insert(0, _ROOT)

router = APIRouter()


class ExportPayload(BaseModel):
    content: str  # Markdown source
    output_path: str | None = None  # Optional explicit output path
    base_dir: str | None = None  # For resolving relative image paths
    dark_mode: bool = False
    font_family: str = "system-ui, sans-serif"
    font_size: int = 14


def _make_output_path(suggested: str | None, suffix: str) -> str:
    if suggested:
        return suggested
    _, path = tempfile.mkstemp(suffix=suffix)
    return path


# ── Endpoints ─────────────────────────────────────────────────────────────────


@router.post("/html")
def export_html(payload: ExportPayload):
    """Export Markdown to a self-contained HTML file, return its path."""
    from backend.renderer import render_markdown

    html = render_markdown(
        payload.content,
        base_dir=payload.base_dir,
        dark_mode=payload.dark_mode,
        font_family=payload.font_family,
        font_size=payload.font_size,
    )
    out_path = _make_output_path(payload.output_path, ".html")
    parent = os.path.dirname(out_path)
    if parent:
        os.makedirs(parent, exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(html)
    return {"path": out_path}


@router.post("/html/download")
def download_html(payload: ExportPayload):
    """Export Markdown to HTML and stream the file for download."""
    from backend.renderer import render_markdown

    html = render_markdown(
        payload.content,
        base_dir=payload.base_dir,
        dark_mode=payload.dark_mode,
        font_family=payload.font_family,
        font_size=payload.font_size,
    )
    _, tmp = tempfile.mkstemp(suffix=".html")
    with open(tmp, "w", encoding="utf-8") as f:
        f.write(html)
    return FileResponse(tmp, media_type="text/html", filename="export.html")


@router.post("/docx")
def export_docx(payload: ExportPayload):
    """Export Markdown to DOCX and return the output file path."""
    import os
    import re
    from io import BytesIO

    import requests
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    out_path = _make_output_path(payload.output_path, ".docx")
    parent = os.path.dirname(out_path)
    if parent:
        os.makedirs(parent, exist_ok=True)

    try:

        def resolve_image_source(src):
            """
            Takes the image src and returns the absolute path to the file.

            :param string src: The src from the image tag.

            :return: Two strings, one indicating whether the image stored locally or online, and another with the path required to access the image.
            """
            src = src.strip().strip("<>")
            if (src.startswith('"') and src.endswith('"')) or (
                src.startswith("'") and src.endswith("'")
            ):
                src = src[1:-1]

            if src.startswith("file://"):
                src = src[7:]

            if re.match(r"^https?://", src, re.IGNORECASE):
                return "remote", src

            if os.path.isabs(src):
                return "local", src

            if payload.base_dir:
                return "local", os.path.abspath(os.path.join(payload.base_dir, src))

            return "local", os.path.abspath(src)

        def insert_image(doc_obj, src):
            """
            Inserts the given image src into the Word document.

            :param Document doc_obj: The Word document to have the image inserted into.
            :param src: The unformatted image source for the image to be inserted.

            :return: A boolean containing true if the operation succeeds, and false if not.
            """
            kind, value = resolve_image_source(src)
            try:
                if kind == "remote":
                    response = requests.get(value, timeout=15, allow_redirects=True)
                    response.raise_for_status()
                    doc_obj.add_picture(BytesIO(response.content), width=Inches(6.2))
                else:
                    if not os.path.exists(value):
                        return False
                    doc_obj.add_picture(value, width=Inches(6.2))

                doc_obj.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                return True
            except Exception:
                return False

        # Create a new Word document
        doc = Document()

        # Parse and convert markdown to Word
        lines = payload.content.split("\n")
        i = 0
        in_code_block = False
        code_block_lines = []

        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.strip().startswith("```"):
                if not in_code_block:
                    in_code_block = True
                    code_block_lines = []
                else:
                    # End of code block
                    in_code_block = False
                    if code_block_lines:
                        code_text = "\n".join(code_block_lines)
                        p = doc.add_paragraph(code_text)
                        p.style = "Intense Quote"
                        for run in p.runs:
                            run.font.name = "Courier New"
                            run.font.size = Pt(10)
                    code_block_lines = []
                i += 1
                continue

            if in_code_block:
                code_block_lines.append(line)
                i += 1
                continue

            # Skip empty lines
            if not line.strip():
                i += 1
                continue

            # Markdown image: ![alt](src)
            md_image_match = re.match(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
            if md_image_match:
                alt_text = md_image_match.group(1).strip()
                src = md_image_match.group(2).strip()
                if not insert_image(doc, src):
                    doc.add_paragraph(alt_text if alt_text else src)
                i += 1
                continue

            # HTML image: <img ... src="..." ...>
            html_image_match = re.match(
                r'^\s*<img\b[^>]*\bsrc=["\']([^"\']+)["\'][^>]*>\s*$',
                line,
                re.IGNORECASE,
            )
            if html_image_match:
                src = html_image_match.group(1).strip()
                alt_match = re.search(r'\balt=["\']([^"\']*)["\']', line, re.IGNORECASE)
                alt_text = alt_match.group(1).strip() if alt_match else ""
                if not insert_image(doc, src):
                    doc.add_paragraph(alt_text if alt_text else src)
                i += 1
                continue

            # Headings
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                text = line.lstrip("#").strip()
                if level == 1:
                    doc.add_heading(text, level=1)
                elif level == 2:
                    doc.add_heading(text, level=2)
                elif level == 3:
                    doc.add_heading(text, level=3)
                else:
                    doc.add_heading(text, level=4)

            # Horizontal rule
            elif line.strip() in ["---", "***", "___"]:
                doc.add_paragraph("_" * 50)

            # Unordered list
            elif re.match(r"^\s*[-*+]\s+", line):
                text = re.sub(r"^\s*[-*+]\s+", "", line)
                doc.add_paragraph(text, style="List Bullet")

            # Ordered list
            elif re.match(r"^\s*\d+\.\s+", line):
                text = re.sub(r"^\s*\d+\.\s+", "", line)
                doc.add_paragraph(text, style="List Number")

            # Blockquote
            elif line.strip().startswith(">"):
                text = line.strip().lstrip(">").strip()
                p = doc.add_paragraph(text, style="Intense Quote")

            # Table detection (simple)
            elif "|" in line and line.strip().startswith("|"):
                table_lines = [line]
                i += 1
                # Collect table rows
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1

                # Parse and create table
                if len(table_lines) > 2:  # Header + separator + at least one row
                    rows_data = []
                    for tline in table_lines:
                        if not re.match(r"^\s*\|[\s:-]+\|", tline):  # Skip separator
                            cells = [c.strip() for c in tline.split("|")[1:-1]]
                            rows_data.append(cells)

                    if rows_data:
                        table = doc.add_table(
                            rows=len(rows_data), cols=len(rows_data[0])
                        )
                        table.style = "Light Grid Accent 1"

                        for row_idx, row_data in enumerate(rows_data):
                            for col_idx, cell_text in enumerate(row_data):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_text
                                if row_idx == 0:  # Header row
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                continue

            # Regular paragraph
            else:
                doc.add_paragraph(line)

            i += 1

        # Save the document
        doc.save(out_path)
        return {"path": out_path}

    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


@router.post("/pdf")
def export_pdf(payload: ExportPayload):
    """Export Markdown to PDF via WeasyPrint and return the output file path."""
    from backend.renderer import render_markdown
    from markdown_reader.plugins.pdf_exporter import export_markdown_to_pdf

    # First render Markdown to HTML
    html = render_markdown(
        payload.content,
        base_dir=payload.base_dir,
        dark_mode=payload.dark_mode,
        font_family=payload.font_family,
        font_size=payload.font_size,
    )

    out_path = _make_output_path(payload.output_path, ".pdf")
    parent = os.path.dirname(out_path)
    if parent:
        os.makedirs(parent, exist_ok=True)
    try:
        export_markdown_to_pdf(html, out_path, base_url=payload.base_dir)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))
    return {"path": out_path}

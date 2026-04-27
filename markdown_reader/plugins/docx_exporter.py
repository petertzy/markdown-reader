from __future__ import annotations

import os
import re
from html import unescape
from html.parser import HTMLParser
from io import BytesIO
from typing import Any

import requests
from docx import Document
from docx.shared import Inches, Pt


def _resolve_image_source(src: str, base_dir: str | None) -> tuple[str, bytes | None]:
    src = src.strip()
    if src.startswith("file://"):
        src = src[7:]
    if re.match(r"^https?://", src, re.IGNORECASE):
        try:
            response = requests.get(src, timeout=15, allow_redirects=True)
            response.raise_for_status()
            return src, response.content
        except Exception:
            return src, None
    if os.path.isabs(src):
        try:
            with open(src, "rb") as f:
                return src, f.read()
        except Exception:
            return src, None
    if base_dir:
        candidate = os.path.abspath(os.path.join(base_dir, src))
        try:
            with open(candidate, "rb") as f:
                return candidate, f.read()
        except Exception:
            return candidate, None
    try:
        with open(src, "rb") as f:
            return src, f.read()
    except Exception:
        return src, None


class _DocxHtmlParser(HTMLParser):
    def __init__(self, document: Document, base_dir: str | None = None):
        super().__init__(convert_charrefs=False)
        self.document = document
        self.base_dir = base_dir
        self.element_stack: list[str] = []
        self.current_paragraph = None
        self.current_table = None
        self.current_row = None
        self.current_cell = None
        self.current_list_style: list[str] = []
        self.current_style: dict[str, Any] = {
            "bold": False,
            "italic": False,
            "underline": False,
            "code": False,
        }
        self.data_buffer = ""
        self.in_pre = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str]]):
        attrs = dict(attrs)
        self._flush_text()
        self.element_stack.append(tag)

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1]) if len(tag) == 2 and tag[1].isdigit() else 1
            self.current_paragraph = self.document.add_heading("", level=level)
        elif tag == "p":
            self.current_paragraph = self.document.add_paragraph()
        elif tag == "blockquote":
            self.current_paragraph = self.document.add_paragraph()
            if self.current_paragraph.paragraph_format is not None:
                self.current_paragraph.paragraph_format.left_indent = Inches(0.25)
        elif tag == "pre":
            self.in_pre = True
            self.current_paragraph = self.document.add_paragraph()
        elif tag == "code":
            self.current_style["code"] = True
        elif tag == "strong" or tag == "b":
            self.current_style["bold"] = True
        elif tag == "em" or tag == "i":
            self.current_style["italic"] = True
        elif tag == "u":
            self.current_style["underline"] = True
        elif tag in ("ul", "ol"):
            style = "List Bullet" if tag == "ul" else "List Number"
            self.current_list_style.append(style)
        elif tag == "li":
            list_style = (
                self.current_list_style[-1] if self.current_list_style else None
            )
            if list_style:
                self.current_paragraph = self.document.add_paragraph(style=list_style)
                level = len(self.current_list_style) - 1
                if level > 0 and self.current_paragraph.paragraph_format is not None:
                    self.current_paragraph.paragraph_format.left_indent = Inches(
                        0.25 * level
                    )
            else:
                self.current_paragraph = self.document.add_paragraph()
        elif tag == "br":
            if self.current_paragraph is None:
                self.current_paragraph = self.document.add_paragraph()
            self.current_paragraph.add_run().add_break()
        elif tag == "img":
            src = attrs.get("src", "").strip()
            if not src:
                return
            image_path, image_data = _resolve_image_source(src, self.base_dir)
            if self.current_paragraph is None:
                self.current_paragraph = self.document.add_paragraph()
            try:
                if image_data is not None:
                    self.current_paragraph.add_run().add_picture(
                        BytesIO(image_data), width=Inches(6.0)
                    )
                else:
                    self.current_paragraph.add_run(f"[{image_path}]")
            except Exception:
                self.current_paragraph.add_run(f"[{image_path}]")
        elif tag == "table":
            self.current_table = []
        elif tag == "tr":
            self.current_row = []
        elif tag in ("td", "th"):
            self.current_cell = []
        elif tag == "a":
            self.current_style["link"] = attrs.get("href")

    def handle_endtag(self, tag: str):
        self._flush_text()
        if not self.element_stack:
            return
        while self.element_stack and self.element_stack[-1] != tag:
            self.element_stack.pop()
        if self.element_stack:
            self.element_stack.pop()

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6", "p", "blockquote", "li", "pre"):
            self.current_paragraph = None
        if tag == "pre":
            self.in_pre = False
        elif tag == "code":
            self.current_style["code"] = False
        elif tag in ("strong", "b"):
            self.current_style["bold"] = False
        elif tag in ("em", "i"):
            self.current_style["italic"] = False
        elif tag == "u":
            self.current_style["underline"] = False
        elif tag == "ul" or tag == "ol":
            if self.current_list_style:
                self.current_list_style.pop()
        elif tag == "table":
            if self.current_table:
                self._build_table()
            self.current_table = None
        elif tag == "tr":
            if self.current_row is not None and self.current_table is not None:
                self.current_table.append(self.current_row)
            self.current_row = None
        elif tag in ("td", "th"):
            if self.current_cell is not None and self.current_row is not None:
                self.current_row.append("".join(self.current_cell).strip())
            self.current_cell = None
        elif tag == "a":
            self.current_style.pop("link", None)

    def handle_data(self, data: str):
        self.data_buffer += data

    def _flush_text(self):
        if not self.data_buffer:
            return
        text = unescape(self.data_buffer)
        self.data_buffer = ""

        if self.current_cell is not None:
            self.current_cell.append(text)
            return

        if self.current_paragraph is None:
            self.current_paragraph = self.document.add_paragraph()

        run = self.current_paragraph.add_run(text)
        run.bold = self.current_style.get("bold", False)
        run.italic = self.current_style.get("italic", False)
        run.underline = self.current_style.get("underline", False)
        if self.current_style.get("code", False):
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        return

    def _build_table(self):
        if not self.current_table:
            return
        num_rows = len(self.current_table)
        num_cols = max(len(row) for row in self.current_table)
        if num_rows == 0 or num_cols == 0:
            return
        table = self.document.add_table(rows=num_rows, cols=num_cols)
        table.style = "Light Grid Accent 1"
        for row_idx, row_data in enumerate(self.current_table):
            for col_idx in range(num_cols):
                cell_text = row_data[col_idx] if col_idx < len(row_data) else ""
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = cell_text
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(11)


def export_html_to_docx(
    html_content: str, output_path: str, base_dir: str | None = None
) -> None:
    """Export rendered HTML into a DOCX file."""
    body_match = re.search(r"<body[^>]*>(.*)</body>", html_content, flags=re.S | re.I)
    body_html = body_match.group(1) if body_match else html_content
    document = Document()
    parser = _DocxHtmlParser(document, base_dir=base_dir)
    parser.feed(body_html)
    parser.close()
    document.save(output_path)

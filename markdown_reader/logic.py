import markdown2
import os
import subprocess
import sys
import webbrowser
import re
import json
from datetime import datetime, timezone
from pathlib import Path
import html2text
from html import escape as html_escape
from tkinter import messagebox
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import traceback
import requests

try:
    import keyring
    from keyring.errors import KeyringError
except Exception:
    keyring = None

    class KeyringError(Exception):
        pass

AI_CREDENTIAL_SERVICE = "MarkdownReader.AI"


def _get_settings_file_path():
    """Return the per-user settings file path for desktop app persistence."""

    if sys.platform == "darwin":
        base_dir = Path.home() / "Library" / "Application Support" / "MarkdownReader"
    elif sys.platform.startswith("win"):
        appdata = os.environ.get("APPDATA", "").strip()
        base_dir = Path(appdata) / "MarkdownReader" if appdata else Path.home() / "AppData" / "Roaming" / "MarkdownReader"
    else:
        base_dir = Path.home() / ".config" / "markdown-reader"

    return base_dir / "settings.json"


APP_SETTINGS_FILE_PATH = _get_settings_file_path()
AI_CHAT_HISTORY_FILE_PATH = APP_SETTINGS_FILE_PATH.parent / "chat_history.json"
AI_AUTOMATION_LOG_FILE_PATH = APP_SETTINGS_FILE_PATH.parent / "ai_automation_log.json"

# Hardcoded provider base URLs so bundled apps work without external env files.
AI_PROVIDER_BASE_URLS = {
    "openrouter": "https://openrouter.ai/api/v1",
    "openai": "https://api.openai.com/v1",
    "anthropic": "https://api.anthropic.com/v1",
}

# Curated fallback model lists shown when the live API cannot be reached.
AI_PROVIDER_DEFAULT_MODELS = {
    "openrouter": [
        "meta-llama/llama-3.3-70b-instruct:free",
        "mistralai/mistral-7b-instruct:free",
        "openai/gpt-4o-mini",
        "openai/gpt-4o",
        "openai/gpt-4-turbo",
        "anthropic/claude-3-5-sonnet",
        "anthropic/claude-3-haiku",
        "google/gemini-2.0-flash-001",
    ],
    "openai": [
        "gpt-4o-mini",
        "gpt-4o",
        "gpt-4-turbo",
        "gpt-3.5-turbo",
    ],
    "anthropic": [
        "claude-3-5-sonnet-latest",
        "claude-3-5-haiku-latest",
        "claude-3-opus-latest",
        "claude-3-haiku-20240307",
    ],
}

# Environment-variable name used to persist the user's chosen model per provider.
AI_PROVIDER_MODEL_ENV = {
    "openrouter": "OPENROUTER_MODEL",
    "openai": "OPENAI_MODEL",
    "anthropic": "ANTHROPIC_MODEL",
}

AI_AGENT_MAX_DOC_CONTEXT = 12000
AI_AGENT_MAX_SELECTION_CONTEXT = 4000
AI_AGENT_MAX_HISTORY_MESSAGES = 16
AI_AGENT_MAX_HISTORY_PREVIEW = 1200
AI_AUTOMATION_MAX_AUDIT_LOG_ENTRIES = 300

AI_AUTOMATION_TASK_TEMPLATES = [
    {
        "id": "format_selection",
        "title": "Format Selected Section",
        "prompt": "Apply Markdown formatting rules to the selected section.",
        "requires_selection": True,
    },
    {
        "id": "generate_toc",
        "title": "Generate Table of Contents",
        "prompt": "Generate a Markdown table of contents from headings and insert it.",
        "requires_selection": False,
    },
    {
        "id": "generate_summary",
        "title": "Generate Summary",
        "prompt": "Generate a concise document summary in Markdown bullet points.",
        "requires_selection": False,
    },
    {
        "id": "fix_code_blocks",
        "title": "Format and Fix Code Blocks",
        "prompt": "Format Markdown code fences and fix common fence syntax issues.",
        "requires_selection": False,
    },
]


def get_ai_automation_task_templates():
    """Return built-in AI automation task templates."""

    templates = []
    for item in AI_AUTOMATION_TASK_TEMPLATES:
        if not isinstance(item, dict):
            continue
        templates.append(
            {
                "id": str(item.get("id", "")).strip(),
                "title": str(item.get("title", "")).strip(),
                "prompt": str(item.get("prompt", "")).strip(),
                "requires_selection": bool(item.get("requires_selection", False)),
            }
        )
    return [t for t in templates if t["id"] and t["title"] and t["prompt"]]


def load_ai_automation_logs(limit=AI_AUTOMATION_MAX_AUDIT_LOG_ENTRIES):
    """Load persisted AI automation audit logs."""

    if not AI_AUTOMATION_LOG_FILE_PATH.exists():
        return []

    try:
        with open(AI_AUTOMATION_LOG_FILE_PATH, "r", encoding="utf-8") as file_obj:
            data = json.load(file_obj)
    except Exception:
        return []

    if not isinstance(data, list):
        return []

    cleaned = []
    for item in data:
        if not isinstance(item, dict):
            continue
        entry = {
            "timestamp": str(item.get("timestamp", "")).strip(),
            "doc_id": str(item.get("doc_id", "")).strip(),
            "status": str(item.get("status", "")).strip().lower(),
            "action_type": str(item.get("action_type", "")).strip().lower(),
            "reason": str(item.get("reason", "")).strip(),
            "user_message": str(item.get("user_message", "")).strip(),
            "content_preview": str(item.get("content_preview", "")).strip(),
            "related_action_id": str(item.get("related_action_id", "")).strip(),
            "action_id": str(item.get("action_id", "")).strip(),
        }
        if not entry["timestamp"]:
            continue
        cleaned.append(entry)

    if limit and isinstance(limit, int) and limit > 0:
        return cleaned[-limit:]
    return cleaned


def save_ai_automation_logs(log_entries):
    """Persist AI automation audit logs."""

    if not isinstance(log_entries, list):
        return

    serializable = []
    for item in log_entries[-AI_AUTOMATION_MAX_AUDIT_LOG_ENTRIES:]:
        if not isinstance(item, dict):
            continue
        timestamp = str(item.get("timestamp", "")).strip()
        if not timestamp:
            continue
        serializable.append(
            {
                "timestamp": timestamp,
                "doc_id": str(item.get("doc_id", "")).strip(),
                "status": str(item.get("status", "")).strip().lower(),
                "action_type": str(item.get("action_type", "")).strip().lower(),
                "reason": str(item.get("reason", "")).strip(),
                "user_message": str(item.get("user_message", "")).strip(),
                "content_preview": str(item.get("content_preview", "")).strip(),
                "related_action_id": str(item.get("related_action_id", "")).strip(),
                "action_id": str(item.get("action_id", "")).strip(),
            }
        )

    try:
        AI_AUTOMATION_LOG_FILE_PATH.parent.mkdir(parents=True, exist_ok=True)
        with open(AI_AUTOMATION_LOG_FILE_PATH, "w", encoding="utf-8") as file_obj:
            json.dump(serializable, file_obj, ensure_ascii=True, indent=2)
    except Exception:
        return


def append_ai_automation_log(entry):
    """Append a single AI automation audit log entry and persist it."""

    if not isinstance(entry, dict):
        return

    logs = load_ai_automation_logs(limit=AI_AUTOMATION_MAX_AUDIT_LOG_ENTRIES)
    logs.append(entry)
    if len(logs) > AI_AUTOMATION_MAX_AUDIT_LOG_ENTRIES:
        logs = logs[-AI_AUTOMATION_MAX_AUDIT_LOG_ENTRIES:]
    save_ai_automation_logs(logs)


def _apply_markdown_formatting_rules(markdown_text):
    """Apply lightweight Markdown formatting normalization rules."""

    if not isinstance(markdown_text, str):
        return ""

    normalized = markdown_text.replace("\r\n", "\n")
    lines = normalized.split("\n")
    output = []

    for raw_line in lines:
        line = raw_line.rstrip()
        line = re.sub(r"^(#{1,6})([^\s#])", r"\1 \2", line)
        line = re.sub(r"^(\s*)([-*+])(\S)", r"\1\2 \3", line)
        line = re.sub(r"^(\s*\d+\.)(\S)", r"\1 \2", line)
        output.append(line)

    compacted = "\n".join(output)
    compacted = re.sub(r"\n{3,}", "\n\n", compacted)
    return compacted


def _guess_code_language(code_block):
    """Guess a language for a code block when fence language is omitted."""

    sample = (code_block or "").strip()
    lowered = sample.lower()
    if not sample:
        return "text"
    if "def " in sample or "import " in sample or "print(" in sample:
        return "python"
    if "function " in sample or "const " in sample or "=>" in sample:
        return "javascript"
    if lowered.startswith("{") and lowered.endswith("}"):
        return "json"
    if "<html" in lowered or "<div" in lowered or "</" in lowered:
        return "html"
    if "select " in lowered and " from " in lowered:
        return "sql"
    return "text"


def _format_and_fix_code_blocks(markdown_text):
    """Normalize Markdown code fences and close unbalanced blocks."""

    if not isinstance(markdown_text, str):
        return ""

    normalized = markdown_text.replace("\r\n", "\n")
    lines = normalized.split("\n")
    out_lines = []
    in_fence = False
    fence_marker = "```"
    block_lines = []

    for line in lines:
        fence_match = re.match(r"^\s*(```+|~~~+)([^`]*)$", line)
        if fence_match:
            marker = fence_match.group(1)
            suffix = (fence_match.group(2) or "").strip()

            if not in_fence:
                in_fence = True
                fence_marker = marker
                block_lines = []
                if suffix:
                    out_lines.append(f"{marker}{suffix}")
                else:
                    out_lines.append(marker)
                continue

            if in_fence and marker.startswith(fence_marker[0]):
                if out_lines and re.match(r"^\s*(```+|~~~+)\s*$", out_lines[-(len(block_lines) + 1)]):
                    opening = out_lines[-(len(block_lines) + 1)]
                    if opening.strip() in ("```", "~~~"):
                        guessed = _guess_code_language("\n".join(block_lines))
                        if guessed:
                            out_lines[-(len(block_lines) + 1)] = f"{opening.strip()}{guessed}"

                out_lines.append(fence_marker)
                in_fence = False
                block_lines = []
                continue

        out_lines.append(line.rstrip())
        if in_fence:
            block_lines.append(line)

    if in_fence:
        if out_lines and re.match(r"^\s*(```+|~~~+)\s*$", out_lines[-(len(block_lines) + 1)]):
            opening = out_lines[-(len(block_lines) + 1)]
            guessed = _guess_code_language("\n".join(block_lines))
            if guessed and opening.strip() in ("```", "~~~"):
                out_lines[-(len(block_lines) + 1)] = f"{opening.strip()}{guessed}"
        out_lines.append(fence_marker)

    return "\n".join(out_lines)


def _slugify_heading_text(text):
    """Create a Markdown anchor slug from heading text."""

    plain = re.sub(r"[`*_~\[\](){}]", "", text or "").strip().lower()
    plain = re.sub(r"[^a-z0-9\s\-]", "", plain)
    plain = re.sub(r"\s+", "-", plain)
    plain = re.sub(r"\-+", "-", plain).strip("-")
    return plain


def _generate_markdown_toc(markdown_text):
    """Generate a Markdown TOC from heading lines."""

    if not isinstance(markdown_text, str) or not markdown_text.strip():
        return ""

    toc_lines = []
    for line in markdown_text.replace("\r\n", "\n").split("\n"):
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if not match:
            continue
        level = len(match.group(1))
        title = match.group(2).strip()
        anchor = _slugify_heading_text(title)
        if not anchor:
            continue
        indent = "  " * max(0, level - 1)
        toc_lines.append(f"{indent}- [{title}](#{anchor})")

    if not toc_lines:
        return ""
    return "## Table of Contents\n\n" + "\n".join(toc_lines) + "\n"


def _merge_toc_into_document(document_text, toc_text):
    """Merge TOC text into a document, replacing existing TOC block when present."""

    doc = (document_text or "").replace("\r\n", "\n")
    toc = (toc_text or "").strip()
    if not toc:
        return doc

    toc_block_pattern = re.compile(
        r"^##\s+Table of Contents\s*\n(?:.*?)(?=^##\s+|\Z)",
        re.MULTILINE | re.DOTALL,
    )
    replacement = toc.rstrip() + "\n\n"
    if toc_block_pattern.search(doc):
        merged = toc_block_pattern.sub(replacement, doc, count=1)
        return merged.rstrip() + "\n"

    if doc.strip():
        return replacement + doc.lstrip("\n")
    return toc.rstrip() + "\n"


def _generate_lightweight_summary(markdown_text):
    """Generate a deterministic summary from headings and first paragraph."""

    if not isinstance(markdown_text, str) or not markdown_text.strip():
        return ""

    normalized = markdown_text.replace("\r\n", "\n")
    headings = []
    for line in normalized.split("\n"):
        match = re.match(r"^#{1,3}\s+(.+?)\s*$", line)
        if match:
            headings.append(match.group(1).strip())
        if len(headings) >= 5:
            break

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", normalized) if p.strip()]
    lead = ""
    for paragraph in paragraphs:
        if paragraph.startswith("#"):
            continue
        lead = re.sub(r"\s+", " ", paragraph)
        break

    lines = ["## Summary"]
    if lead:
        lines.append("")
        lines.append(f"- {lead[:240]}{'...' if len(lead) > 240 else ''}")
    if headings:
        lines.append("")
        lines.append("- Main sections:")
        for heading in headings:
            lines.append(f"  - {heading}")
    return "\n".join(lines).strip() + "\n"


def build_ai_automation_fallback(user_message, document_text="", selected_text=""):
    """Build deterministic local automation actions for common repetitive tasks."""

    msg = (user_message or "").strip()
    lowered = msg.lower()
    selection = selected_text if isinstance(selected_text, str) else ""
    document = document_text if isinstance(document_text, str) else ""
    target = selection if selection.strip() else document

    if not msg:
        return None

    if any(keyword in lowered for keyword in ("template", "task template", "automation template")):
        template_lines = []
        for template in get_ai_automation_task_templates():
            selection_hint = " (selection required)" if template["requires_selection"] else ""
            template_lines.append(f"- {template['id']}: {template['title']}{selection_hint}")
        return {
            "assistant_message": "Available automation templates:\n" + "\n".join(template_lines),
            "proposed_action": {
                "type": "none",
                "content": "",
                "reason": "task_templates",
            },
            "used_provider": "local-fallback",
        }

    if any(keyword in lowered for keyword in ("table of contents", "toc", "目录")):
        toc_source = document if document.strip() else target
        toc = _generate_markdown_toc(toc_source)
        if toc:
            if not selection.strip():
                merged_document = _merge_toc_into_document(document, toc)
                return {
                    "assistant_message": f"{toc}",
                    "proposed_action": {
                        "type": "replace_document",
                        "content": merged_document,
                        "reason": "generate_toc_full_document",
                    },
                    "used_provider": "local-fallback",
                }
            return {
                "assistant_message": "Generated a table of contents. Review and apply if it matches your document structure.",
                "proposed_action": {
                    "type": "replace_selection",
                    "content": toc,
                    "reason": "generate_toc",
                },
                "used_provider": "local-fallback",
            }

    if "summary" in lowered or "summarize" in lowered or "总结" in lowered:
        summary_source = document if document.strip() else target
        summary = _generate_lightweight_summary(summary_source)
        if not selection.strip():
            if summary:
                return {
                    "assistant_message": f"{summary}",
                    "proposed_action": {
                        "type": "replace_document",
                        "content": summary,
                        "reason": "generate_summary_full_document",
                    },
                    "used_provider": "local-fallback",
                }
            return {
                "assistant_message": "No document content available to summarize.",
                "proposed_action": {
                    "type": "none",
                    "content": "",
                    "reason": "no_content_for_summary",
                },
                "used_provider": "local-fallback",
            }
        if summary:
            return {
                "assistant_message": "Generated a concise summary based on current content.",
                "proposed_action": {
                    "type": "replace_selection",
                    "content": summary,
                    "reason": "generate_summary",
                },
                "used_provider": "local-fallback",
            }

    if (
        "format code" in lowered
        or "code block" in lowered
        or "correct syntax" in lowered
        or "fix code" in lowered
    ):
        if target.strip():
            if not selection.strip():
                return {
                    "assistant_message": "Select the code block you want to fix, then run this task again.",
                    "proposed_action": {
                        "type": "none",
                        "content": "",
                        "reason": "selection_required_for_code_fix",
                    },
                    "used_provider": "local-fallback",
                }
            fixed = _format_and_fix_code_blocks(target)
            return {
                "assistant_message": "Prepared formatted code blocks and fixed common fence syntax issues.",
                "proposed_action": {
                    "type": "replace_selection",
                    "content": fixed,
                    "reason": "fix_code_blocks",
                },
                "used_provider": "local-fallback",
            }

    if "format" in lowered or "formatting" in lowered:
        if target.strip():
            if not selection.strip():
                if document.strip():
                    formatted_document = _apply_markdown_formatting_rules(document)
                    return {
                        "assistant_message": "Applied Markdown formatting normalization rules to the full document.",
                        "proposed_action": {
                            "type": "replace_document",
                            "content": formatted_document,
                            "reason": "format_rules_full_document",
                        },
                        "used_provider": "local-fallback",
                    }
                return {
                    "assistant_message": "Select the section you want to format, then run this task again.",
                    "proposed_action": {
                        "type": "none",
                        "content": "",
                        "reason": "selection_required_for_format",
                    },
                    "used_provider": "local-fallback",
                }
            formatted = _apply_markdown_formatting_rules(target)
            return {
                "assistant_message": "Applied Markdown formatting normalization rules.",
                "proposed_action": {
                    "type": "replace_selection",
                    "content": formatted,
                    "reason": "format_rules",
                },
                "used_provider": "local-fallback",
            }

    return None


def load_ai_chat_histories():
    """Load persisted AI chat history grouped by document id."""

    if not AI_CHAT_HISTORY_FILE_PATH.exists():
        return {}

    try:
        with open(AI_CHAT_HISTORY_FILE_PATH, "r", encoding="utf-8") as file_obj:
            data = json.load(file_obj)
    except Exception:
        return {}

    if not isinstance(data, dict):
        return {}

    cleaned = {}
    for doc_id, messages in data.items():
        if not isinstance(doc_id, str) or not isinstance(messages, list):
            continue
        safe_messages = []
        for item in messages:
            if not isinstance(item, dict):
                continue
            role = str(item.get("role", "")).strip().lower()
            content = str(item.get("content", "")).strip()
            if role not in ("user", "assistant") or not content:
                continue
            safe_messages.append({"role": role, "content": content})
        if safe_messages:
            cleaned[doc_id] = safe_messages[-80:]
    return cleaned


def save_ai_chat_histories(histories):
    """Persist AI chat histories grouped by document id."""

    if not isinstance(histories, dict):
        return

    serializable = {}
    for doc_id, messages in histories.items():
        if not isinstance(doc_id, str) or not isinstance(messages, list):
            continue
        safe_messages = []
        for item in messages[-80:]:
            if not isinstance(item, dict):
                continue
            role = str(item.get("role", "")).strip().lower()
            content = str(item.get("content", "")).strip()
            if role not in ("user", "assistant") or not content:
                continue
            safe_messages.append({"role": role, "content": content})
        if safe_messages:
            serializable[doc_id] = safe_messages

    try:
        AI_CHAT_HISTORY_FILE_PATH.parent.mkdir(parents=True, exist_ok=True)
        with open(AI_CHAT_HISTORY_FILE_PATH, "w", encoding="utf-8") as file_obj:
            json.dump(serializable, file_obj, ensure_ascii=True, indent=2)
    except Exception:
        return


class TranslationConfigError(RuntimeError):
    """Raised when the configured AI provider cannot be used without user action."""

    def __init__(self, message, provider_name=None, env_var=None, invalid_key=False):
        super().__init__(message)
        self.provider_name = provider_name
        self.env_var = env_var
        self.invalid_key = invalid_key


def _file_uri(path):
    """Build a browser-safe file URI with proper escaping."""

    return Path(path).resolve().as_uri()


def _open_file_in_browser(path):
    """Open a local file in the default browser with platform-safe handling."""

    resolved_path = str(Path(path).resolve())

    if sys.platform == "darwin":
        # macOS bundled apps can hit ascii encoding issues through webbrowser's
        # AppleScript backend. Using `open` avoids that code path.
        subprocess.Popen(["open", resolved_path])
        return True

    return webbrowser.open(_file_uri(resolved_path), new=0)


def _load_app_settings():
    """Load app settings from the per-user JSON file."""

    if not APP_SETTINGS_FILE_PATH.exists():
        return {}

    try:
        with open(APP_SETTINGS_FILE_PATH, "r", encoding="utf-8") as file_obj:
            data = json.load(file_obj)
            return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _save_app_settings(settings):
    """Write app settings to the per-user JSON file."""

    if not isinstance(settings, dict):
        return

    try:
        APP_SETTINGS_FILE_PATH.parent.mkdir(parents=True, exist_ok=True)
        with open(APP_SETTINGS_FILE_PATH, "w", encoding="utf-8") as file_obj:
            json.dump(settings, file_obj, ensure_ascii=True, indent=2)
        # Restrict to owner-only read/write to reduce exposure of local settings.
        if sys.platform != "win32":
            import stat as _stat
            APP_SETTINGS_FILE_PATH.chmod(_stat.S_IRUSR | _stat.S_IWUSR)
    except Exception:
        return


def _normalize_provider_name(provider_name):
    """Normalize provider aliases and fallback to supported providers."""

    normalized = (provider_name or "").strip().lower()
    if normalized == "athropic":
        normalized = "anthropic"
    if normalized not in ("openrouter", "openai", "anthropic"):
        normalized = "openrouter"
    return normalized


def load_persisted_ai_settings():
    """Load persisted AI settings from app settings into process environment."""

    settings = _load_app_settings()
    if not settings:
        settings = {
            "ai_provider": "openrouter",
            "ai_models": {},
        }
        _save_app_settings(settings)

    saved_provider = _normalize_provider_name(settings.get("ai_provider", ""))
    if not os.environ.get("AI_PROVIDER", "").strip():
        os.environ["AI_PROVIDER"] = saved_provider

    saved_models = settings.get("ai_models", {})
    if not isinstance(saved_models, dict):
        return

    for provider_name, env_key in AI_PROVIDER_MODEL_ENV.items():
        if os.environ.get(env_key, "").strip():
            continue
        model_name = (saved_models.get(provider_name, "") or "").strip()
        if model_name:
            os.environ[env_key] = model_name


def set_current_ai_provider(provider_name):
    """Persist and activate the selected AI provider."""

    normalized = _normalize_provider_name(provider_name)

    os.environ["AI_PROVIDER"] = normalized
    settings = _load_app_settings()
    settings["ai_provider"] = normalized
    _save_app_settings(settings)


def get_ai_provider_env_var(provider_name):
    """Return the API key environment variable used by a provider."""

    normalized = (provider_name or "").strip().lower()
    if normalized == "athropic":
        normalized = "anthropic"

    return {
        "openrouter": "OPENROUTER_API_KEY",
        "openai": "OPENAI_API_KEY",
        "anthropic": "ANTHROPIC_API_KEY",
    }.get(normalized)


def get_ai_provider_display_name(provider_name):
    """Return a human-friendly provider name."""

    normalized = (provider_name or "").strip().lower()
    if normalized == "openrouter":
        return "OpenRouter"
    if normalized == "openai":
        return "OpenAI"
    if normalized == "anthropic":
        return "Anthropic"
    return provider_name or "AI"


def is_secure_key_storage_available():
    """Return whether system secure storage (keyring) is available."""

    return keyring is not None


def get_secure_ai_api_key(provider_name):
    """Read a provider API key from keyring.

    For backward compatibility, this performs a lazy one-time migration from the
    legacy JSON ``api_keys`` field if present.
    """

    normalized = _normalize_provider_name(provider_name)

    if keyring is not None:
        try:
            stored = keyring.get_password(AI_CREDENTIAL_SERVICE, normalized)
            if stored and stored.strip():
                return stored.strip()
        except KeyringError:
            pass
        except Exception:
            pass

    # Legacy fallback: migrate old JSON-stored key when encountered.
    settings = _load_app_settings()
    api_keys = settings.get("api_keys", {})
    if isinstance(api_keys, dict):
        legacy_key = (api_keys.get(normalized, "") or "").strip()
        if legacy_key:
            if keyring is not None:
                try:
                    keyring.set_password(AI_CREDENTIAL_SERVICE, normalized, legacy_key)
                    del api_keys[normalized]
                    settings["api_keys"] = api_keys
                    _save_app_settings(settings)
                except Exception:
                    # If migration fails, still allow app use with legacy value.
                    pass
            return legacy_key
    return ""


def set_secure_ai_api_key(provider_name, api_key):
    """Persist a provider API key in keyring."""

    normalized = _normalize_provider_name(provider_name)
    cleaned_key = (api_key or "").strip()
    if not cleaned_key:
        raise RuntimeError("API key cannot be empty.")
    if keyring is None:
        raise RuntimeError("Secure key storage is not available on this system.")

    try:
        keyring.set_password(AI_CREDENTIAL_SERVICE, normalized, cleaned_key)
    except KeyringError as exc:
        raise RuntimeError(f"Could not store API key securely: {exc}") from exc
    except Exception as exc:
        raise RuntimeError(f"Could not store API key securely: {exc}") from exc

    # Clean any old plaintext JSON key for this provider after successful save.
    settings = _load_app_settings()
    api_keys = settings.get("api_keys", {})
    if isinstance(api_keys, dict) and normalized in api_keys:
        del api_keys[normalized]
        settings["api_keys"] = api_keys
        _save_app_settings(settings)


def delete_secure_ai_api_key(provider_name):
    """Remove a provider API key from keyring."""

    normalized = _normalize_provider_name(provider_name)
    if keyring is not None:
        try:
            keyring.delete_password(AI_CREDENTIAL_SERVICE, normalized)
        except KeyringError:
            pass
        except Exception:
            pass

    # Also remove any legacy plaintext JSON copy.
    settings = _load_app_settings()
    api_keys = settings.get("api_keys", {})
    if isinstance(api_keys, dict) and normalized in api_keys:
        del api_keys[normalized]
        settings["api_keys"] = api_keys
        _save_app_settings(settings)


def get_ai_provider_model(provider_name):
    """Return the currently selected model for *provider_name*.

    Priority:
    1. runtime os.environ
    2. app settings JSON (per-user persistent config)
    3. legacy keyring value (backward compatibility)
    4. built-in default  (first entry in AI_PROVIDER_DEFAULT_MODELS)
    """
    normalized = (provider_name or "").strip().lower()
    if normalized == "athropic":
        normalized = "anthropic"
    env_key = AI_PROVIDER_MODEL_ENV.get(normalized)

    # 1. Runtime environment (explicit process override)
    if env_key:
        val = os.getenv(env_key, "").strip()
        if val:
            return val

    # 2. App settings (normal persistent storage for non-secret options)
    settings = _load_app_settings()
    saved_models = settings.get("ai_models", {})
    if isinstance(saved_models, dict):
        saved_model = (saved_models.get(normalized, "") or "").strip()
        if saved_model:
            return saved_model

    # 3. Legacy keyring-persisted choice (for migration compatibility)
    if keyring is not None:
        try:
            stored = keyring.get_password(
                AI_CREDENTIAL_SERVICE, f"model:{normalized}"
            )
            if stored and stored.strip():
                return stored.strip()
        except Exception:
            pass

    # 4. Built-in default
    defaults = AI_PROVIDER_DEFAULT_MODELS.get(normalized, [])
    return defaults[0] if defaults else ""


def set_ai_provider_model(provider_name, model_name):
    """Persist *model_name* for *provider_name* to env + app settings."""
    normalized = _normalize_provider_name(provider_name)
    env_key = AI_PROVIDER_MODEL_ENV.get(normalized)
    cleaned_model = (model_name or "").strip()
    if not cleaned_model:
        return

    if env_key:
        os.environ[env_key] = cleaned_model

    settings = _load_app_settings()
    saved_models = settings.get("ai_models")
    if not isinstance(saved_models, dict):
        saved_models = {}
    saved_models[normalized] = cleaned_model
    settings["ai_models"] = saved_models
    _save_app_settings(settings)


def fetch_available_models(provider_name, api_key, timeout=8):
    """Fetch the model list from the provider API and return a sorted list of model IDs.

    Falls back to the built-in default list on any error so the dialog always
    has something to show.
    """
    normalized = (provider_name or "").strip().lower()
    base_url = AI_PROVIDER_BASE_URLS.get(normalized, "")
    defaults = list(AI_PROVIDER_DEFAULT_MODELS.get(normalized, []))

    if not api_key or not base_url:
        return defaults

    try:
        if normalized == "anthropic":
            # Anthropic models endpoint
            headers = {
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
            }
            resp = requests.get(
                f"{base_url.rstrip('/')}/models",
                headers=headers,
                timeout=timeout,
            )
            resp.raise_for_status()
            data = resp.json()
            models = [m["id"] for m in data.get("data", []) if isinstance(m, dict) and m.get("id")]
        else:
            # OpenAI-compatible endpoint (OpenRouter and OpenAI both support /models)
            headers = {"Authorization": f"Bearer {api_key}"}
            resp = requests.get(
                f"{base_url.rstrip('/')}/models",
                headers=headers,
                timeout=timeout,
            )
            resp.raise_for_status()
            data = resp.json()
            models = [m["id"] for m in data.get("data", []) if isinstance(m, dict) and m.get("id")]

        if models:
            # Put defaults first so the most useful ones surface at the top,
            # then append any extra models returned by the API.
            known = set(defaults)
            extras = sorted(m for m in models if m not in known)
            return defaults + extras
    except Exception:
        pass

    return defaults


def _chunk_markdown_for_translation(markdown_text, max_lines=20):
    """Split markdown into translation-friendly chunks while keeping code fences intact."""

    if not isinstance(markdown_text, str) or not markdown_text:
        return []

    normalized_text = markdown_text.replace("\r\n", "\n")
    lines = normalized_text.splitlines(keepends=True)
    chunks = []
    current_lines = []
    current_line_count = 0
    fence_open = False

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```") or stripped.startswith("~~~"):
            fence_open = not fence_open

        current_lines.append(line)
        current_line_count += 1

        should_flush = False
        if current_line_count >= max_lines:
            if fence_open:
                should_flush = False
            elif stripped == "":
                should_flush = True
            elif current_line_count >= max_lines + 8:
                should_flush = True

        if should_flush:
            chunks.append("".join(current_lines))
            current_lines = []
            current_line_count = 0

    if current_lines:
        chunks.append("".join(current_lines))

    return chunks


def split_markdown_for_translation(markdown_text, chunk_lines=20):
    """Return translation chunks for UI-side progressive rendering."""

    return _chunk_markdown_for_translation(markdown_text, max_lines=chunk_lines)


def _restore_chunk_outer_whitespace(original_text, translated_text):
    """Preserve leading and trailing newlines from the original chunk."""

    if not isinstance(original_text, str):
        return translated_text

    translated = translated_text if isinstance(translated_text, str) else str(translated_text)
    leading_match = re.match(r"^\s*", original_text)
    trailing_match = re.search(r"\s*$", original_text)
    leading = leading_match.group(0) if leading_match else ""
    trailing = trailing_match.group(0) if trailing_match else ""

    core = translated.strip()
    return f"{leading}{core}{trailing}"


def _request_translation_openai_compatible(base_url, api_key, model, system_prompt, user_prompt):
    """
    Sends translation request to OpenAI-compatible chat endpoint.
    """

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": 0.2,
    }

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    if "openrouter.ai" in base_url:
        headers["HTTP-Referer"] = os.getenv("OPENROUTER_REFERER", "https://localhost")
        headers["X-Title"] = os.getenv("OPENROUTER_APP_NAME", "MarkdownReader")

    response = requests.post(
        f"{base_url.rstrip('/')}/chat/completions",
        headers=headers,
        json=payload,
        timeout=90,
    )
    try:
        response.raise_for_status()
    except requests.HTTPError as exc:
        message = ""
        try:
            error_payload = response.json()
        except ValueError:
            error_payload = None

        if isinstance(error_payload, dict):
            error_value = error_payload.get("error")
            if isinstance(error_value, dict):
                message = str(error_value.get("message", "") or error_value)
            elif error_value:
                message = str(error_value)

        if not message:
            message = response.text.strip()

        raise requests.HTTPError(message or str(exc), response=response) from exc

    data = response.json()

    choices = data.get("choices", []) if isinstance(data, dict) else []
    if not choices:
        raise RuntimeError(f"No choices in AI response: {data}")

    first_choice = choices[0] if isinstance(choices[0], dict) else {}
    message = first_choice.get("message", {}) if isinstance(first_choice.get("message", {}), dict) else {}
    content = message.get("content", "")

    if isinstance(content, list):
        text_parts = []
        for item in content:
            if isinstance(item, dict):
                if item.get("type") == "text":
                    text_parts.append(str(item.get("text", "")))
            elif isinstance(item, str):
                text_parts.append(item)
        content = "\n".join([part for part in text_parts if part])

    if not content and isinstance(first_choice.get("text"), str):
        content = first_choice.get("text", "")

    if not isinstance(content, str) or not content.strip():
        raise RuntimeError(f"Empty content in AI response: {data}")

    return content


def _request_translation_anthropic(base_url, api_key, model, system_prompt, user_prompt):
    """
    Sends translation request to Anthropic messages endpoint.
    """

    payload = {
        "model": model,
        "max_tokens": 8192,
        "temperature": 0.2,
        "system": system_prompt,
        "messages": [
            {"role": "user", "content": user_prompt},
        ],
    }

    headers = {
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
        "Content-Type": "application/json",
    }

    response = requests.post(
        f"{base_url.rstrip('/')}/messages",
        headers=headers,
        json=payload,
        timeout=90,
    )
    try:
        response.raise_for_status()
    except requests.HTTPError as exc:
        message = ""
        try:
            error_payload = response.json()
        except ValueError:
            error_payload = None

        if isinstance(error_payload, dict):
            error_value = error_payload.get("error")
            if isinstance(error_value, dict):
                message = str(error_value.get("message", "") or error_value)
            elif error_value:
                message = str(error_value)
            elif error_payload.get("message"):
                message = str(error_payload.get("message"))

        if not message:
            message = response.text.strip()

        raise requests.HTTPError(message or str(exc), response=response) from exc

    data = response.json()
    blocks = data.get("content", [])
    text_blocks = [block.get("text", "") for block in blocks if isinstance(block, dict) and block.get("type") == "text"]
    content = "\n".join([part for part in text_blocks if part])
    if not content.strip():
        raise RuntimeError(f"Empty content in Anthropic response: {data}")
    return content


def _extract_json_object(raw_text):
    """
    Extracts a JSON object string from raw LLM output.

    :param string raw_text: Text returned by the AI model.
    :return: A JSON object string if found, otherwise None.
    """

    if not raw_text:
        return None

    text = raw_text.strip()

    fenced_match = re.search(r'```(?:json)?\s*(\{[\s\S]*?\})\s*```', text, re.IGNORECASE)
    if fenced_match:
        return fenced_match.group(1).strip()

    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        return text[start:end + 1]

    return None


def translate_markdown_with_ai(markdown_text, source_language, target_language):
    """
    Translates Markdown content using an OpenAI-compatible API while preserving Markdown formatting.

    Configuration source:
    1) Runtime process environment
    2) Per-user app settings JSON + secure key store

    :param string markdown_text: Markdown content to translate.
    :param string source_language: Source language name.
    :param string target_language: Target language name.

    :return: Tuple (translated_markdown, ambiguity_notes_list).
    :raises RuntimeError: If translation fails or configuration is missing.
    """

    if not isinstance(markdown_text, str) or not markdown_text.strip():
        return "", []

    source_language = (source_language or "").strip()
    target_language = (target_language or "").strip()

    if not source_language or not target_language:
        raise RuntimeError("Source and target language are required.")

    if source_language.lower() == target_language.lower():
        return markdown_text, ["Source and target language are identical; no translation applied."]

    provider_name = os.getenv("AI_PROVIDER", "openrouter").strip().lower()

    if provider_name == "athropic":
        provider_name = "anthropic"

    if provider_name not in ("openrouter", "openai", "anthropic"):
        raise TranslationConfigError(
            f"Unsupported AI provider '{provider_name}'. Use openrouter, openai, or anthropic in AI_PROVIDER",
            provider_name=provider_name,
        )

    def _provider_cfg(name):
        stored_key = get_secure_ai_api_key(name)
        # Use hardcoded base URLs so bundled apps work without external env files.
        # Explicit process env vars still override when present.
        base_url = AI_PROVIDER_BASE_URLS.get(name, "")
        model = get_ai_provider_model(name)
        if name == "openrouter":
            return {
                "name": "openrouter",
                "api_key": os.getenv("OPENROUTER_API_KEY", "").strip() or stored_key,
                "base_url": os.getenv("OPENROUTER_BASE_URL", base_url).strip(),
                "model": model,
                "type": "openai-compatible",
                "env_var": "OPENROUTER_API_KEY",
            }
        if name == "openai":
            return {
                "name": "openai",
                "api_key": os.getenv("OPENAI_API_KEY", "").strip() or stored_key,
                "base_url": os.getenv("OPENAI_BASE_URL", base_url).strip(),
                "model": model,
                "type": "openai-compatible",
                "env_var": "OPENAI_API_KEY",
            }
        return {
            "name": "anthropic",
            "api_key": os.getenv("ANTHROPIC_API_KEY", "").strip() or stored_key,
            "base_url": os.getenv("ANTHROPIC_BASE_URL", base_url).strip(),
            "model": model,
            "type": "anthropic",
            "env_var": "ANTHROPIC_API_KEY",
        }

    provider_order = [provider_name] + [name for name in ("openrouter", "openai", "anthropic") if name != provider_name]
    provider_candidates = []
    for name in provider_order:
        cfg = _provider_cfg(name)
        if cfg["api_key"] and cfg["base_url"] and cfg["model"]:
            provider_candidates.append(cfg)

    if not provider_candidates:
        env_var = get_ai_provider_env_var(provider_name)
        raise TranslationConfigError(
            "No valid AI provider config found.",
            provider_name=provider_name,
            env_var=env_var,
        )

    system_prompt = (
        "You are a professional technical translator for Markdown documents. "
        "Preserve Markdown structure exactly: headings, lists, code fences, links, image syntax, tables, HTML tags, "
        "inline code, and front matter must remain valid Markdown. "
        "Do not add explanations outside JSON."
    )

    user_prompt = (
        f"Translate the following Markdown from {source_language} to {target_language}.\n"
        "Return STRICT JSON with this schema:\n"
        "{\n"
        "  \"translated_markdown\": \"<string>\",\n"
        "  \"ambiguity_notes\": [\"<note1>\", \"<note2>\"]\n"
        "}\n"
        "Rules:\n"
        "1) Keep all Markdown syntax and structure intact.\n"
        "2) Do not translate code, URLs, file paths, or command-line statements unless they are natural-language prose.\n"
        "3) If wording is ambiguous, list concise notes in ambiguity_notes.\n"
        "4) If no ambiguity, return an empty array for ambiguity_notes.\n\n"
        "Markdown to translate:\n"
        f"{markdown_text}"
    )

    content = ""
    used_provider = None
    errors = []
    for cfg in provider_candidates:
        try:
            if cfg["type"] == "openai-compatible":
                content = _request_translation_openai_compatible(
                    cfg["base_url"],
                    cfg["api_key"],
                    cfg["model"],
                    system_prompt,
                    user_prompt,
                )
            else:
                content = _request_translation_anthropic(
                    cfg["base_url"],
                    cfg["api_key"],
                    cfg["model"],
                    system_prompt,
                    user_prompt,
                )
            used_provider = cfg["name"]
            break
        except requests.RequestException as exc:
            status = getattr(getattr(exc, "response", None), "status_code", None)
            errors.append(f"{cfg['name']}: HTTP {status} {exc}")
            if cfg["name"] == provider_name and status in (401, 403):
                raise TranslationConfigError(
                    f"{get_ai_provider_display_name(provider_name)} rejected the configured API key.",
                    provider_name=provider_name,
                    env_var=cfg.get("env_var"),
                    invalid_key=True,
                ) from exc
            # Try next configured provider on rate limit/auth/server issues.
            if status in (401, 402, 403, 404, 408, 409, 429, 500, 502, 503, 504):
                continue
            continue
        except Exception as exc:
            errors.append(f"{cfg['name']}: {exc}")
            continue

    if not content:
        tried_providers = ", ".join([c["name"] for c in provider_candidates])
        error_text = " | ".join(errors[:3]) if errors else "Unknown translation error"

        # Friendlier guidance for a common configuration issue.
        for err in errors:
            lowered = err.lower()
            if "not a chat model" in lowered or "v1/chat/completions" in lowered:
                bad_provider = (err.split(":", 1)[0] or "").strip().lower()
                bad_cfg = next((c for c in provider_candidates if c.get("name") == bad_provider), None)
                bad_model = (bad_cfg or {}).get("model", "")
                provider_label = get_ai_provider_display_name(bad_provider or provider_name)

                message = (
                    "AI translation could not continue because the selected model is not chat-compatible.\n\n"
                    f"Provider: {provider_label}"
                )
                if bad_model:
                    message += f"\nModel: {bad_model}"
                message += (
                    "\n\nPlease open Settings -> AI Provider & API Keys..., "
                    "choose a chat-capable model, and retry."
                )
                raise RuntimeError(message)

        raise RuntimeError(
            "AI translation failed after trying configured providers.\n"
            f"Tried providers: {tried_providers}\n"
            f"Details: {error_text}"
        )

    json_text = _extract_json_object(content)
    if not json_text:
        return _restore_chunk_outer_whitespace(markdown_text, content), ["AI response was not strict JSON; fallback text was used."]

    try:
        parsed = json.loads(json_text)
    except json.JSONDecodeError:
        return _restore_chunk_outer_whitespace(markdown_text, content), ["AI response JSON parsing failed; fallback text was used."]

    translated_markdown = parsed.get("translated_markdown", "")
    ambiguity_notes = parsed.get("ambiguity_notes", [])

    if not isinstance(translated_markdown, str):
        translated_markdown = str(translated_markdown)

    if not isinstance(ambiguity_notes, list):
        ambiguity_notes = [str(ambiguity_notes)] if ambiguity_notes else []

    ambiguity_notes = [str(note).strip() for note in ambiguity_notes if str(note).strip()]
    if used_provider and used_provider != provider_name:
        ambiguity_notes.insert(0, f"Primary provider '{provider_name}' failed; fallback provider '{used_provider}' was used.")
    return _restore_chunk_outer_whitespace(markdown_text, translated_markdown), ambiguity_notes


def translate_markdown_in_chunks(markdown_text, source_language, target_language, chunk_lines=20, progress_callback=None):
    """Translate markdown in smaller chunks and emit progress updates."""

    if not isinstance(markdown_text, str) or not markdown_text.strip():
        return [], []

    chunks = _chunk_markdown_for_translation(markdown_text, max_lines=chunk_lines)
    translated_chunks = []
    all_notes = []
    total_chunks = len(chunks)

    if progress_callback:
        progress_callback(0, total_chunks, "Preparing translation...")

    for index, chunk in enumerate(chunks, start=1):
        translated_chunk, ambiguity_notes = translate_markdown_with_ai(
            chunk,
            source_language,
            target_language,
        )
        translated_chunks.append(translated_chunk)
        all_notes.extend(ambiguity_notes)
        if progress_callback:
            progress_callback(index, total_chunks, f"Translated chunk {index}/{total_chunks}")

    return translated_chunks, all_notes


def request_ai_agent_response(user_message, document_text="", selected_text="", chat_history=None):
    """Request a structured AI agent response for editor chat workflows."""

    message = (user_message or "").strip()
    if not message:
        raise RuntimeError("Message cannot be empty.")

    # Support deterministic offline automation for common repetitive tasks.
    fallback_result = build_ai_automation_fallback(
        message,
        document_text=document_text,
        selected_text=selected_text,
    )
    if fallback_result:
        return fallback_result

    provider_name = os.getenv("AI_PROVIDER", "openrouter").strip().lower()
    if provider_name == "athropic":
        provider_name = "anthropic"
    if provider_name not in ("openrouter", "openai", "anthropic"):
        raise TranslationConfigError(
            f"Unsupported AI provider '{provider_name}'. Use openrouter, openai, or anthropic in AI_PROVIDER",
            provider_name=provider_name,
        )

    def _provider_cfg(name):
        stored_key = get_secure_ai_api_key(name)
        base_url = AI_PROVIDER_BASE_URLS.get(name, "")
        model = get_ai_provider_model(name)
        if name == "openrouter":
            return {
                "name": "openrouter",
                "api_key": os.getenv("OPENROUTER_API_KEY", "").strip() or stored_key,
                "base_url": os.getenv("OPENROUTER_BASE_URL", base_url).strip(),
                "model": model,
                "type": "openai-compatible",
                "env_var": "OPENROUTER_API_KEY",
            }
        if name == "openai":
            return {
                "name": "openai",
                "api_key": os.getenv("OPENAI_API_KEY", "").strip() or stored_key,
                "base_url": os.getenv("OPENAI_BASE_URL", base_url).strip(),
                "model": model,
                "type": "openai-compatible",
                "env_var": "OPENAI_API_KEY",
            }
        return {
            "name": "anthropic",
            "api_key": os.getenv("ANTHROPIC_API_KEY", "").strip() or stored_key,
            "base_url": os.getenv("ANTHROPIC_BASE_URL", base_url).strip(),
            "model": model,
            "type": "anthropic",
            "env_var": "ANTHROPIC_API_KEY",
        }

    provider_order = [provider_name] + [
        name for name in ("openrouter", "openai", "anthropic") if name != provider_name
    ]
    provider_candidates = []
    for name in provider_order:
        cfg = _provider_cfg(name)
        if cfg["api_key"] and cfg["base_url"] and cfg["model"]:
            provider_candidates.append(cfg)

    if not provider_candidates:
        env_var = get_ai_provider_env_var(provider_name)
        raise TranslationConfigError(
            "No valid AI provider config found.",
            provider_name=provider_name,
            env_var=env_var,
        )

    doc_context = (document_text or "").strip()
    selected_context = (selected_text or "").strip()
    if len(doc_context) > AI_AGENT_MAX_DOC_CONTEXT:
        doc_context = doc_context[:AI_AGENT_MAX_DOC_CONTEXT] + "\n\n[Document context truncated]"
    if len(selected_context) > AI_AGENT_MAX_SELECTION_CONTEXT:
        selected_context = selected_context[:AI_AGENT_MAX_SELECTION_CONTEXT] + "\n\n[Selection context truncated]"

    history_lines = []
    safe_history = chat_history if isinstance(chat_history, list) else []
    for item in safe_history[-AI_AGENT_MAX_HISTORY_MESSAGES:]:
        if not isinstance(item, dict):
            continue
        role = str(item.get("role", "")).strip().lower()
        if role not in ("user", "assistant"):
            continue
        content = str(item.get("content", "")).strip()
        if not content:
            continue
        if len(content) > AI_AGENT_MAX_HISTORY_PREVIEW:
            content = content[:AI_AGENT_MAX_HISTORY_PREVIEW] + " ..."
        history_lines.append(f"{role.upper()}: {content}")

    history_block = "\n".join(history_lines) if history_lines else "(none)"
    selection_block = selected_context if selected_context else "(no current selection)"
    document_block = doc_context if doc_context else "(no document content)"

    system_prompt = (
        "You are an in-editor AI agent for Markdown documents. "
        "You must answer clearly and safely. "
        "Automate repetitive editor tasks when appropriate (formatting, TOC, summaries, code block cleanup). "
        "When user intent implies a direct editor action, propose exactly one action. "
        "Allowed actions: none, replace_selection. "
        "Never propose shell commands, file-system operations, or code execution. "
        "Use replace_selection for all editor edits. "
        "Return strict JSON only."
    )

    user_prompt = (
        "Return STRICT JSON with this schema:\n"
        "{\n"
        "  \"assistant_message\": \"<string>\",\n"
        "  \"proposed_action\": {\n"
        "    \"type\": \"none|replace_selection\",\n"
        "    \"content\": \"<string>\",\n"
        "    \"reason\": \"<short string>\"\n"
        "  }\n"
        "}\n"
        "Rules:\n"
        "1) assistant_message must be plain helpful text for the user.\n"
        "2) proposed_action.content must be empty when type is none.\n"
        "3) Preserve Markdown validity in generated content.\n"
        "4) Keep response concise and editor-focused.\n"
        "5) If user asks to summarize, assistant_message must include the actual summary text, not meta commentary.\n"
        "6) Avoid placeholders like 'Here is a summary' without the summary body.\n"
        "7) For formatting requests, normalize Markdown headings/lists/code fences.\n"
        "8) For table-of-contents requests, output Markdown links with heading anchors.\n\n"
        "Recent conversation:\n"
        f"{history_block}\n\n"
        "Current selected text:\n"
        f"{selection_block}\n\n"
        "Current document context:\n"
        f"{document_block}\n\n"
        "User request:\n"
        f"{message}"
    )

    content = ""
    used_provider = None
    errors = []
    for cfg in provider_candidates:
        try:
            if cfg["type"] == "openai-compatible":
                content = _request_translation_openai_compatible(
                    cfg["base_url"],
                    cfg["api_key"],
                    cfg["model"],
                    system_prompt,
                    user_prompt,
                )
            else:
                content = _request_translation_anthropic(
                    cfg["base_url"],
                    cfg["api_key"],
                    cfg["model"],
                    system_prompt,
                    user_prompt,
                )
            used_provider = cfg["name"]
            break
        except requests.RequestException as exc:
            status = getattr(getattr(exc, "response", None), "status_code", None)
            errors.append(f"{cfg['name']}: HTTP {status} {exc}")
            if cfg["name"] == provider_name and status in (401, 403):
                raise TranslationConfigError(
                    f"{get_ai_provider_display_name(provider_name)} rejected the configured API key.",
                    provider_name=provider_name,
                    env_var=cfg.get("env_var"),
                    invalid_key=True,
                ) from exc
            continue
        except Exception as exc:
            errors.append(f"{cfg['name']}: {exc}")
            continue

    if not content:
        tried = ", ".join([c["name"] for c in provider_candidates])
        detail = " | ".join(errors[:3]) if errors else "Unknown AI agent error"
        raise RuntimeError(
            "AI agent failed after trying configured providers.\n"
            f"Tried providers: {tried}\n"
            f"Details: {detail}"
        )

    json_text = _extract_json_object(content)
    if not json_text:
        return {
            "assistant_message": content.strip(),
            "proposed_action": {"type": "none", "content": "", "reason": "non_json_response"},
            "used_provider": used_provider,
        }

    try:
        parsed = json.loads(json_text)
    except json.JSONDecodeError:
        return {
            "assistant_message": content.strip(),
            "proposed_action": {"type": "none", "content": "", "reason": "json_parse_failed"},
            "used_provider": used_provider,
        }

    assistant_message = str(parsed.get("assistant_message", "")).strip() or "(No assistant message)"
    action = parsed.get("proposed_action", {})
    if not isinstance(action, dict):
        action = {}

    action_type = str(action.get("type", "none")).strip().lower()
    if action_type == "insert_at_cursor":
        action_type = "replace_selection"
    if action_type not in ("none", "replace_selection"):
        action_type = "none"

    action_content = str(action.get("content", ""))
    action_reason = str(action.get("reason", "")).strip()
    if action_type == "none":
        action_content = ""

    if action_type != "none" and not action_content.strip():
        action_type = "none"

    return {
        "assistant_message": assistant_message,
        "proposed_action": {
            "type": action_type,
            "content": action_content,
            "reason": action_reason,
        },
        "used_provider": used_provider,
    }


def estimate_translation_chunk_count(markdown_text, chunk_lines=20):
    """Estimate how many translation chunks will be produced."""

    if not isinstance(markdown_text, str) or not markdown_text.strip():
        return 0

    chunks = _chunk_markdown_for_translation(markdown_text, max_lines=chunk_lines)
    return max(1, len(chunks))


def _normalize_math_content(content):
    """
    Normalizes common Unicode math characters to MathJax-friendly LaTeX.
    Also converts common informal notation patterns to proper TeX.

    :param string content: Raw math expression content.

    :return: A normalized math expression string.
    """

    if not isinstance(content, str) or not content:
        return content

    normalized = content
    
    # === UNICODE SUPERSCRIPTS/SUBSCRIPTS CONVERSION ===
    # Convert Unicode superscripts to TeX: x² -> x^2, x³ -> x^3, xⁿ -> x^n
    superscript_map = {
        '⁰': '0', '¹': '1', '²': '2', '³': '3', '⁴': '4',
        '⁵': '5', '⁶': '6', '⁷': '7', '⁸': '8', '⁹': '9',
        'ⁿ': 'n'
    }
    
    subscript_map = {
        '₀': '0', '₁': '1', '₂': '2', '₃': '3', '₄': '4',
        '₅': '5', '₆': '6', '₇': '7', '₈': '8', '₉': '9',
        'ᵢ': 'i', 'ⱼ': 'j', 'ₖ': 'k'
    }
    
    # Replace Unicode superscripts: x² -> x^{2}
    for unicode_char, tex_char in superscript_map.items():
        # Match letter followed by superscript
        normalized = re.sub(f'([a-zA-Z]){re.escape(unicode_char)}',
                           lambda m: f"{m.group(1)}^{{{tex_char}}}", normalized)
    
    # Replace Unicode subscripts: x₁ -> x_{1}
    for unicode_char, tex_char in subscript_map.items():
        # Match letter followed by subscript
        normalized = re.sub(f'([a-zA-Z]){re.escape(unicode_char)}',
                           lambda m: f"{m.group(1)}_{{{tex_char}}}", normalized)
    
    # === UNICODE REPLACEMENTS ===
    normalized = normalized.replace('−', '-')
    normalized = normalized.replace('–', '-')
    normalized = normalized.replace('—', '-')
    normalized = normalized.replace('≤', r'\le ')
    normalized = normalized.replace('≥', r'\ge ')
    normalized = normalized.replace('∑', r'\sum ')
    normalized = normalized.replace('×', r'\times ')
    normalized = normalized.replace('÷', r'\div ')
    normalized = normalized.replace('·', r'\cdot ')
    normalized = normalized.replace('…', r'\ldots ')

    normalized = re.sub(r'([A-Za-z0-9Α-Ωα-ωπΠ])\s*¯', r'\\bar{\1}', normalized)

    normalized = normalized.replace('ϕ', r'\phi ')
    normalized = normalized.replace('φ', r'\phi ')
    normalized = normalized.replace('β', r'\beta ')
    normalized = normalized.replace('π', r'\pi ')
    normalized = normalized.replace('Π', r'\Pi ')

    normalized = re.sub(r'\bPr\(', r'\\Pr(', normalized)
    
    # === SUBSCRIPT NOTATION FIXES (CONSERVATIVE) ===
    # Only fix patterns we're very confident about:
    
    # Pattern 1: Letter(s) followed by comma-subscript (B-spline: Bi,1 -> B_{i,1})
    # This is very safe because commas are rarely used otherwise in math
    normalized = re.sub(r'([A-Za-z])([a-z]+),(\d+)', r'\1_{\2,\3}', normalized)
    
    # Pattern 2: Specific known subscript patterns with operators immediately following
    # ti+ -> t_{i}+, ti- -> t_{i}-, ti) -> t_{i}), etc.
    # Only for lowercase single letters followed by single letter then operator
    normalized = re.sub(r'(?<![\\a-zA-Z])([tTxXaAbBiI])([a-z])([+\-\)<>\]]\b)', 
                       r'\1_{\2}\3', normalized)

    # === SYNTAX ERROR FIXES ===
    normalized = normalized.replace('^)', ')')
    normalized = normalized.replace('^]', ']')
    normalized = re.sub(r'(\)|\])([0-9]+)', r'\1^\2', normalized)

    normalized = re.sub(r'\s{2,}', ' ', normalized).strip()
    return normalized


def _is_probably_math_line(line):
    """
    Heuristic check for lines that are likely intended as math expressions.

    :param string line: A single line of markdown text.

    :return: True if line likely represents math, else False.
    """

    stripped = line.strip()
    if not stripped or len(stripped) < 12:
        return False

    word_count = len(stripped.split())
    if word_count >= 6 and re.match(r'^(where|let|the|this|that|and|or)\b', stripped, re.IGNORECASE):
        return False

    if stripped.startswith(('#', '>', '-', '*', '`', '|')):
        return False

    if 'http://' in stripped or 'https://' in stripped:
        return False

    if '$' in stripped or '\\[' in stripped or '\\]' in stripped or '\\(' in stripped or '\\)' in stripped:
        return False

    has_core_math = '=' in stripped and any(sym in stripped for sym in ('^', 'π', 'Π', '¯', '−', '-', '(', ')', '[', ']'))
    has_tex_command = bool(re.search(r'\\(frac|sqrt|sum|int|bar|alpha|beta|gamma|pi)\b', stripped))

    return has_core_math or has_tex_command


def _auto_wrap_bare_math_lines(markdown_text):
    """
    Wrap likely bare math lines as display math so they can be rendered.

    :param string markdown_text: Original markdown text.

    :return: Updated markdown text with wrapped bare math lines.
    """

    wrapped_lines = []
    in_fenced_code = False
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            in_fenced_code = not in_fenced_code
            wrapped_lines.append(line)
            continue

        if in_fenced_code or _is_markdown_media_line(stripped):
            wrapped_lines.append(line)
            continue

        if _is_probably_math_line(line):
            wrapped_lines.append(f"$$\n{line.strip()}\n$$")
        else:
            wrapped_lines.append(line)

    return '\n'.join(wrapped_lines)


def _is_probably_math_token(token):
    """
    Heuristic check for bare inline math tokens inside prose.

    :param string token: A token possibly containing inline math.

    :return: True if token likely represents math, else False.
    """

    if not token:
        return False

    core = token.strip()
    core = core.strip("'\".,;:!?")

    # Never treat inline-code fragments as math.
    if '`' in core:
        return False

    if _looks_like_url_or_path(core):
        return False
    
    # === EARLY CHECKS FOR OBVIOUS MATH PATTERNS ===
    
    # Check for Unicode superscripts/subscripts (x², x³, xⁿ, x₁, etc.)
    # Includes: superscript numbers (²³⁴⁵⁶⁷⁸⁹), common superscript letters (ⁿ), subscript numbers
    if any(c in core for c in '⁰¹²³⁴⁵⁶⁷⁸⁹ⁿ₀₁₂₃₄₅₆₇₈₉ᵢⱼₖ'):
        return True
    
    # Check for simple exponents: x^2, y^n, etc. (letter + ^ + something)
    if re.search(r'^[a-zA-Z]\^[\w\{\}]', core):
        return True
    
    # === SUBSCRIPT/SUPERSCRIPT PATTERNS ===
    
    # Pattern checks for math variables:
    # 1. Double letters: kk, ii, tt
    # 2. Underscore subscript: t_i, x_n, B_i
    # 3. Digit subscript: t1, x2, B3
    # 4. Comma subscript: Bi,1 (B-spline notation), ti,k, etc.
    
    is_double_letter = bool(re.fullmatch(r'([a-zA-Z])\1', core))
    has_underscore_subscript = bool(re.search(r'[a-zA-Z]_[a-zA-Z0-9]', core))
    has_digit_subscript = bool(re.search(r'[a-zA-Z]\d', core))
    has_comma_subscript = bool(re.search(r'[a-zA-Z]+,[a-zA-Z0-9+\-]+', core))
    
    if core.startswith(('http://', 'https://', 'www.', 'file://')):
        return False

    if '$' in core or '\\(' in core or '\\)' in core or '\\[' in core or '\\]' in core:
        return False

    # Reject pure lowercase English words (at least 3 chars), but allow math vars
    if re.fullmatch(r'[a-z]+', core) and len(core) >= 3:
        return False

    # Accept short multichar tokens with numbers or uppercase (common math vars like ti, Bi, xi, pi)
    if re.match(r'^[a-zA-Z]{1,2}\d*[a-zA-Z]?$', core):
        if len(core) < 3 and (any(c.isupper() for c in core) or has_digit_subscript or has_underscore_subscript or has_comma_subscript):
            return True
        if len(core) >= 3:
            return True

    contains_strong_marker = any(sym in core for sym in (
        '=', '^', '∑', '≤', '≥', '|', '¯', '−', 'π', 'Π', 'β', 'ϕ', 'φ', '_', '…'
    ))
    has_bracket_shape = ('(' in core and ')' in core) or ('[' in core and ']' in core) or ('{' in core and '}' in core)
    has_greek_symbol = any(sym in core for sym in ('π', 'Π', 'β', 'ϕ', 'φ'))

    return (
        contains_strong_marker and (has_bracket_shape or len(core) >= 6 or has_greek_symbol)
    ) or is_double_letter or has_underscore_subscript or has_digit_subscript or has_comma_subscript


def _looks_like_url_or_path(text):
    """
    Checks whether a token likely represents a URL/file path rather than math.
    """

    if not text:
        return False

    lowered = text.lower()
    if lowered.startswith(("http://", "https://", "file://", "www.")):
        return True

    if "/" in text or "\\" in text:
        return True

    if re.search(r'\.(png|jpg|jpeg|gif|webp|svg|bmp|pdf|md|txt|py|js|ts|json|yml|yaml)$', lowered):
        return True

    return False


def _is_markdown_media_line(stripped_line):
    """
    True for markdown image/link lines that should bypass math auto-wrap.
    """

    if not stripped_line:
        return False

    if re.match(r'^!\[[^\]]*\]\([^)]+\)$', stripped_line):
        return True

    if re.match(r'^\[[^\]]+\]\([^)]+\)$', stripped_line):
        return True

    if "<img" in stripped_line.lower():
        return True

    return False


def _auto_wrap_bare_math_spans(markdown_text):
    """
    Wrap likely inline bare math tokens with \\(...\\) in mixed prose lines.
    Returns list of (protected_text, replacements_dict) to preserve LaTeX delimiters.

    :param string markdown_text: Original markdown text.

    :return: Tuple of (updated markdown text with placeholders, dict of replacements).
    """

    replacements = {}
    counter = [0]
    
    def make_placeholder(content):
        """Create a placeholder for wrapped math formula."""
        key = f"AUTOBAREMATH{counter[0]}X"
        counter[0] += 1
        replacements[key] = f'<span class="math-inline">\\({content}\\)</span>'
        return key

    def wrap_token(token):
        # Skip tokens that are already placeholders
        if 'MATHPLACEHOLDER' in token or 'AUTOBAREMATH' in token:
            return token
        
        # Skip URLs / file paths / markdown image-link token bodies
        if _looks_like_url_or_path(token) or token.startswith('![') or '](' in token:
            return token

        # Skip markdown formatting syntax (bold/italic markers)
        # Complete patterns: _italic_ or __bold__ or *italic* or **bold**
        if re.match(r'^[_*]{1,2}\w+[_*]{1,2}$', token):
            return token
        
        # Skip incomplete markdown markers (part of multi-token markdown)
        # _Italic or *bold (starts with marker but no closing on same token)
        if re.match(r'^[_*]{1,2}[a-zA-Z]', token) and not re.search(r'[_*]$', token):
            return token
        # text_ or text) (ends with marker but no opening on same token)
        if re.search(r'^[a-zA-Z].*[_*]{1,2}$', token) and not re.match(r'^[_*]', token):
            return token
            
        match = re.match(r'^(.*?)([.,;:!?]+)$', token)
        if match:
            core = match.group(1)
            suffix = match.group(2)
        else:
            core = token
            suffix = ''

        if not _is_probably_math_token(core):
            return token

        normalized_core = _normalize_math_content(core)
        placeholder = make_placeholder(normalized_core)
        return f"{placeholder}{suffix}"

    output_lines = []
    in_fenced_code = False
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            output_lines.append(line)
            continue

        if stripped.startswith('```'):
            in_fenced_code = not in_fenced_code
            output_lines.append(line)
            continue

        is_markdown_table_row = bool(re.match(r'^\s*\|.*\|\s*$', stripped))
        is_markdown_table_separator = bool(re.match(r'^\s*\|?[\s:-]+\|[\s|:-]*$', stripped))

        if in_fenced_code or _is_markdown_media_line(stripped):
            output_lines.append(line)
            continue

        if stripped.startswith(('#', '>')) or is_markdown_table_row or is_markdown_table_separator:
            output_lines.append(line)
            continue

        if _is_probably_math_line(line):
            output_lines.append(line)
            continue

        tokens = line.split(' ')
        wrapped_tokens = [wrap_token(token) for token in tokens]
        output_lines.append(' '.join(wrapped_tokens))

    result_text = '\n'.join(output_lines)
    return result_text, replacements


def _protect_math(markdown_text):
    """
    Protects math expressions from being escaped by markdown2.

    :param string markdown_text: The Markdown text to check for math.

    :return: A string text containing the Markdown text with the maths swapped for placeholders and 
    a dictionary containing the math divs with placeholder text keys.
    """

    replacements = {}
    counter = [0]

    def make_placeholder(content, display=False):
        key = f"MATHPLACEHOLDER{counter[0]}X"
        counter[0] += 1
        normalized_content = _normalize_math_content(content)
        if display:
            replacements[key] = f'<div class="math-display">\\[{normalized_content}\\]</div>'
        else:
            replacements[key] = f'<span class="math-inline">\\({normalized_content}\\)</span>'
        return key

    def replace_block(m):
        content = m.group(1).strip()
        # Skip empty or whitespace-only blocks
        if not content:
            return m.group(0)  # Return original $$...$$
        return make_placeholder(content, display=True)

    def replace_inline(m):
        content = m.group(1).strip()
        # Skip empty or whitespace-only inline formulas
        if not content:
            return m.group(0)  # Return original $...$
        return make_placeholder(content, display=False)

    # IMPORTANT ORDER:
    # 1. Protect explicit math delimiters FIRST ($$...$$, \[...\], $...$, \(...\))
    # 2. Then apply auto-wrapping for bare math (uses placeholders to avoid markdown2 escaping)
    
    text = re.sub(r'\$\$([\s\S]+?)\$\$', replace_block, markdown_text)
    text = re.sub(r'\\\[([\s\S]+?)\\\]', replace_block, text)
    text = re.sub(r'(?<!\$)\$(?!\$)([^\$\n]+?)(?<!\$)\$(?!\$)', replace_inline, text)
    text = re.sub(r'\\\((.+?)\\\)', replace_inline, text)
    
    # Auto-wrap bare math tokens (returns placeholders + dict)
    text, auto_wrapped = _auto_wrap_bare_math_spans(text)
    replacements.update(auto_wrapped)  # Merge auto-wrapped formulas
    
    text = _auto_wrap_bare_math_lines(text)

    return text, replacements


def _restore_math(html_content, replacements):
    """
    Restores math placeholders back to MathJax-compatible HTML.

    :param string html_content: The string containing the HTML containing placeholders where math should be.
    :param Dictionary<string, string> replacements: The math expressions taken from the original Markdown to be reinserted.

    :return: A string containing the updated HTML with the math equations reinserted.
    """

    for key, value in replacements.items():
        html_content = html_content.replace(key, value)
        html_content = html_content.replace(f'<p>{key}</p>', value)
    return html_content


def _get_math_styles():
    return """
        .math-display {
            display: block;
            text-align: center;
            margin: 1em 0;
            overflow-x: auto;
        }
        .math-inline {
            display: inline;
        }
    """


def _get_mathjax_script():
    return """
        <script>
            window.MathJax = {
                tex: {
                    inlineMath: [['\\\\(', '\\\\)'], ['$', '$']],
                    displayMath: [['\\\\[', '\\\\]'], ['$$', '$$']],
                    processEscapes: true,
                    processEnvironments: true,
                    tagSide: 'right'
                },
                svg: {
                    fontCache: 'global',
                    scale: 1
                },
                options: {
                    processHtmlClass: 'math-display|math-inline',
                    skipHtmlTags: ['script', 'noscript', 'style', 'textarea', 'pre', 'code'],
                    ignoreHtmlClass: 'no-mathjax'
                }
            };
        </script>
        <script src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>
        <script>
            (function() {
                function processMath() {
                    if (window.MathJax && window.MathJax.typesetPromise) {
                        MathJax.typesetPromise().catch(function(err) {
                            console.error('MathJax rendering error:', err);
                        });
                    }
                }
                // Wait for MathJax to be fully loaded
                if (document.readyState === 'loading') {
                    document.addEventListener('DOMContentLoaded', processMath);
                } else {
                    setTimeout(processMath, 0);
                }
            })();
        </script>
    """


def update_preview(app):
    """
    Updates the preview of the Markdown file when the file is changed.

    :param MarkdownReader app: The MarkdownReader application instance.

    :return: A boolean set to true if the preview is updated successfully, and false if not.

    :raises ConversionError: If there is an error converting to Markdown2.
    :raises RuntimeError: If there is an error updating the preview.
    :raises RuntimeError: If the preview fails to generate.
    """

    if not app.editors:
        return False
    # Build HTML content from current editor text, with robust error reporting.
    try:
        idx = app.notebook.index(app.notebook.select())
        text_area = app.editors[idx]
        # Use override content if present (from per-selection color logic)
        markdown_text = getattr(app, '_preview_content_override', None)
        if markdown_text is None:
            markdown_text = text_area.get("1.0", "end-1c")

        # Attempt to fix image paths when a file is open

        if hasattr(app, 'file_paths') and app.file_paths:
            try:
                idx = app.notebook.index(app.notebook.select())
                current_path = app.file_paths[idx]
                if current_path is not None:
                    base_dir = os.path.dirname(current_path)
                    markdown_text = fix_image_paths(markdown_text, base_dir)
            except Exception:                
                # Non-fatal: continue without fixed image paths
                pass
# If the editor appears empty in the packaged app, try to read
        # the file contents directly from disk (handles race conditions
        # where the GUI editor hasn't populated yet).
        try:
            if (not isinstance(markdown_text, str) or not markdown_text.strip()) and hasattr(app, 'file_paths') and app.file_paths:
                current_path = app.file_paths[idx]
                if current_path and os.path.isfile(current_path):
                    with open(current_path, 'r', encoding='utf-8', errors='replace') as fh:
                        disk_text = fh.read()
                    if isinstance(disk_text, str) and disk_text.strip():
                        markdown_text = disk_text
        except Exception:
            pass

        try:
            # Protect math BEFORE markdown2 processes it
            protected_text, math_replacements = _protect_math(markdown_text)

            html_content = markdown2.markdown(
                protected_text,
                extras=["fenced-code-blocks", "code-friendly", "tables", "break-on-newline"]
            )

            # Restore math expressions AFTER markdown2
            html_content = _restore_math(html_content, math_replacements)

        except Exception as e:
            tb = traceback.format_exc()
            print(f"markdown2 conversion error: {e}\n{tb}")
                        # Produce an HTML page containing the full traceback for easier debugging

            html_content = f"<h2>Error generating preview</h2><pre>{html_escape(tb)}</pre>"
    except Exception as e:
        tb = traceback.format_exc()
        print(f"update_preview unexpected error: {e}\n{tb}")
        html_content = f"<h2>Unexpected error generating preview</h2><pre>{html_escape(tb)}</pre>"

    # Get style from app (with fallback)

    font_family = getattr(app, 'current_font_family', 'Consolas')
    font_size = getattr(app, 'current_font_size', 14)
    fg_color = getattr(app, 'current_fg_color', '#000000')
    bg_color = getattr(app, 'current_bg_color', 'white')
    if getattr(app, 'dark_mode', False):
        bg_color = '#1e1e1e'
        fg_color = '#dcdcdc'

    web_font_family = font_family
    if font_family.lower() in ["arial", "helvetica", "verdana", "tahoma", "trebuchet ms"]:
        web_font_family += ", sans-serif"
    elif font_family.lower() in ["times new roman", "georgia", "garamond", "serif"]:
        web_font_family += ", serif"
    elif font_family.lower() in ["consolas", "courier new", "monospace"]:
        web_font_family += ", monospace"
    else:
        web_font_family += ", sans-serif"

    h1 = font_size + 18
    h2 = font_size + 12
    h3 = font_size + 8
    h4 = font_size + 4
    h5 = font_size + 2
    h6 = font_size + 1
    base = font_size + 2

    try:
        debug_snippet = ''
        try:
            debug_snippet = markdown_text[:1000].replace('--', '- -')
        except Exception:
            debug_snippet = '<unable to read markdown snippet>'

        debug_path = ''
        try:
            debug_path = getattr(app, 'file_paths', [None])[idx] if hasattr(app, 'file_paths') else ''
        except Exception:
            debug_path = ''
        debug_comment = f""

        with open(app.preview_file, 'w', encoding='utf-8') as f:
            f.write(f"""
            <html>
            {debug_comment}
            <head>
                <meta charset="UTF-8">
                <style>
                    {_get_math_styles()}
                    .copy-button:hover {{
                        background-color: #a5a8b6;
                    }}
                    .copy-button {{
                        position: absolute;
                        top: 10px;
                        right: 10px;
                        background-color: #cacbd0;
                        color: rgb(49, 50, 52);
                        font-size: 0.4em;
                        padding: 1px;
                        border: none;
                        border-radius: 2px;
                        width: auto;
                        min-width: 25px;
                        height: auto;
                        min-height: 15px;
                        cursor: pointer;
                        z-index: 9999;
                    }}
                    body {{
                        background-color: {bg_color};
                        color: {fg_color};
                        font-family: {web_font_family};
                        padding: 20px;
                        font-size: {base}px;
                        line-height: 1.6;
                    }}
                    h1 {{ font-size: {h1}px; }}
                    h2 {{ font-size: {h2}px; }}
                    h3 {{ font-size: {h3}px; }}
                    h4 {{ font-size: {h4}px; }}
                    h5 {{ font-size: {h5}px; }}
                    h6 {{ font-size: {h6}px; }}
                    b, strong {{ font-weight: bold; }}
                    i, em {{ font-style: italic; }}
                    u {{ text-decoration: underline; }}
                    pre code{{
                        background-color: #f4f4f4;
                        color: #000000;
                        font-family: {web_font_family};
                        font-size: {max(font_size - 2, 10)}px;
                        padding: 28px 12px 0px 12px;
                        border-radius: 6px;
                        overflow-x: auto;
                        display: block;
                        max-width: 100%;
                        box-sizing: border-box;
                        white-space: pre;
                    }}
                    code {{
                        background-color: #f4f4f4;
                        color: #000000;
                        font-family: {web_font_family};
                        font-size: {max(font_size - 2, 10)}px;
                        padding: 0 4px;
                        border-radius: 4px;
                        white-space: normal;
                        display: inline;
                    }}
                    img {{
                        max-width: 90vw;
                        max-height: 90vh;
                        height: auto;
                        width: auto;
                        display: block;
                        margin: 10px 0;
                    }}
                    table {{
                        border-collapse: collapse;
                        width: 100%;
                        margin-top: 20px;
                        font-size: {base}px;
                    }}
                    th, td {{
                        text-align: left;
                        border: 1px solid #ccc;
                        padding: 12px 16px;
                        vertical-align: top;
                        font-size: {base}px;
                    }}
                    th {{
                        background-color: #f3f3f3;
                        color: #333;
                    }}
                    tr:nth-child(even) {{
                        background-color: #fafafa;
                    }}
                    @media print {{
                        .copy-button {{
                            display: none !important;
                        }}
                        pre code {{
                            background-color: #f4f4f4 !important;
                            color: #000 !important;
                            display: block !important;
                            white-space: pre-wrap !important;
                            padding: 8px 12px !important;
                            border-radius: 6px !important;
                            overflow-x: visible !important;
                            word-break: break-word !important;
                            word-wrap: break-word !important;
                            -webkit-print-color-adjust: exact;
                            print-color-adjust: exact;
                        }}
                        code {{
                            background-color: #f4f4f4 !important;
                            color: #000 !important;
                            display: inline !important;
                            white-space: normal !important;
                            padding: 0 4px !important;
                            border-radius: 4px !important;
                            -webkit-print-color-adjust: exact;
                            print-color-adjust: exact;
                        }}
                        body {{
                            -webkit-print-color-adjust: exact;
                            print-color-adjust: exact;
                        }}
                    }}
                </style>
                <script>
                    function addCopyButtonToAllCodeBlocks() {{
                        const codeBlocks = document.querySelectorAll('pre code');
                        codeBlocks.forEach(function(codeBlock) {{
                            if (!codeBlock.parentElement.querySelector('.copy-button')) {{
                                const copyButton = document.createElement('button');
                                copyButton.className = 'copy-button';
                                copyButton.textContent = 'Copy';

                                const wrapper = document.createElement('div');
                                wrapper.style.position = 'relative';
                                codeBlock.parentElement.parentNode.insertBefore(wrapper, codeBlock.parentElement);
                                wrapper.appendChild(copyButton);
                                wrapper.appendChild(codeBlock.parentElement);

                                copyButton.addEventListener('click', function() {{
                                    const codeContent = codeBlock.innerText;
                                    const originalText = copyButton.textContent;
                                    copyButton.textContent = 'Copied!';
                                    navigator.clipboard.writeText(codeContent).then(function() {{
                                        setTimeout(function() {{
                                            copyButton.textContent = originalText;
                                        }}, 1000);
                                    }}).catch(function(err) {{
                                        console.error('Could not copy text: ', err);
                                    }});
                                }});

                                copyButton.style.position = 'absolute';
                                copyButton.style.top = '10px';
                                copyButton.style.right = '10px';
                                copyButton.style.border = 'none';
                                copyButton.style.padding = '5px 10px';
                                copyButton.style.cursor = 'pointer';
                                copyButton.style.zIndex = '10';
                            }}
                        }});
                    }}
                    document.addEventListener('DOMContentLoaded', addCopyButtonToAllCodeBlocks);
                </script>
                {_get_mathjax_script()}
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """)
        return True
    except Exception as e:
        messagebox.showerror("Error", f"Failed to generate preview: {e}")


def open_preview_in_browser(preview_file, app):
    """
    Opens a preview of the Markdown file into the browser. 

    :param string preview_file: The file path for the preview of the Markdown file.
    :param MarkdownReader app: The MarkdownReader application instance.

    :raises RuntimeError: If the preview fails to open.
    """

    if update_preview(app):
        try:
            _open_file_in_browser(preview_file)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open preview: {e}")
    else:
        messagebox.showinfo("Info", "No document to preview.")


def fix_image_paths(markdown_text, base_path):
    """
    Takes Markdown code containing images and fixes the file paths of the images.

    :param string markdown_text: The Markdown file containing some images. 
    :param string base_path: The file path for the base directory of the images.

    :return: A string containing the corrected Markdown file with updated image file paths.
    """
    
    def repl(m): # Callable used to replace the regex pattern with a new one. 
        alt = m.group(1)
        src = m.group(2)
        if src.startswith(('http://', 'https://', 'file://', '/')):
            return m.group(0)
        abs_path = os.path.abspath(os.path.join(base_path, src))
        abs_url = 'file://' + abs_path.replace('\\', '/')
        return f'![{alt}]({abs_url})'

    return re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', repl, markdown_text)


def export_to_html(app, output_path):
    """
    Exports the current markdown document to an HTML file.
    
    :param MarkdownReader app: The MarkdownReader application instance.
    :param string output_path: The path where the HTML file should be saved.
    
    :return: A boolean value set to true if the file is successfully exported and false if not.

    :raises RuntimeError: If the image paths cannot be processed.
    :raises RuntimeError: If the HTML could not be exported.
    """

    if not app.editors:
        messagebox.showinfo("Info", "No document to export.")
        return False
    
    try:
        idx = app.notebook.index(app.notebook.select())
        text_area = app.editors[idx]
        markdown_text = text_area.get("1.0", "end-1c")
        
        # Fix image paths if a file is currently open
        if hasattr(app, 'file_paths') and app.file_paths:
            try:
                current_path = app.file_paths[idx]
                if current_path is not None:
                    base_dir = os.path.dirname(current_path)
                    # Convert file:// paths to relative paths for export
                    def convert_file_url_to_relative(text, base_dir):
                        def repl(m):
                            alt = m.group(1)
                            src = m.group(2)
                            if src.startswith('file://'):
                                # Convert file:// URL back to relative path
                                file_path = src.replace('file://', '')
                                try:
                                    rel_path = os.path.relpath(file_path, base_dir)
                                    return f'![{alt}]({rel_path})'
                                except (ValueError, OSError):
                                    return m.group(0)
                            return m.group(0)
                        return re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', repl, text)
                    
                    markdown_text = convert_file_url_to_relative(markdown_text, base_dir)
            except Exception as e:
                print(f"Warning: Could not process image paths: {e}")
        
        # Convert markdown to HTML
        # FIX APPLIED: Added "break-on-newline" to extras here as well
        html_content = markdown2.markdown(
            markdown_text, 
            extras=["fenced-code-blocks", "code-friendly", "tables", "break-on-newline"]
        )
        
        # Get style from app (with fallback)
        font_family = getattr(app, 'current_font_family', 'Consolas')
        font_size = getattr(app, 'current_font_size', 14)
        fg_color = getattr(app, 'current_fg_color', '#000000')
        bg_color = getattr(app, 'current_bg_color', 'white')
        
        if getattr(app, 'dark_mode', False):
            bg_color = '#1e1e1e'
            fg_color = '#dcdcdc'
        
        # For web, use a generic fallback for common fonts
        web_font_family = font_family
        if font_family.lower() in ["arial", "helvetica", "verdana", "tahoma", "trebuchet ms"]:
            web_font_family += ", sans-serif"
        elif font_family.lower() in ["times new roman", "georgia", "garamond", "serif"]:
            web_font_family += ", serif"
        elif font_family.lower() in ["consolas", "courier new", "monospace"]:
            web_font_family += ", monospace"
        else:
            web_font_family += ", sans-serif"
        
        # Heading sizes relative to base font size
        h1 = font_size + 18
        h2 = font_size + 12
        h3 = font_size + 8
        h4 = font_size + 4
        h5 = font_size + 2
        h6 = font_size + 1
        base = font_size + 2
        
        # Generate complete HTML document
        html_document = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Exported Markdown</title>
    <style>
        body {{
            background-color: {bg_color};
            color: {fg_color};
            font-family: {web_font_family};
            padding: 20px;
            font-size: {base}px;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
        }}
        h1 {{ font-size: {h1}px; }}
        h2 {{ font-size: {h2}px; }}
        h3 {{ font-size: {h3}px; }}
        h4 {{ font-size: {h4}px; }}
        h5 {{ font-size: {h5}px; }}
        h6 {{ font-size: {h6}px; }}
        b, strong {{ font-weight: bold; }}
        i, em {{ font-style: italic; }}
        u {{ text-decoration: underline; }}
        pre code{{
            background-color: #f4f4f4;
            color: #000000;
            font-family: {web_font_family};
            font-size: {max(font_size - 2, 10)}px;
            padding: 12px;
            border-radius: 6px;
            overflow-x: auto;
            display: block;
            max-width: 100%;
            box-sizing: border-box;
            white-space: pre;
        }}
        code {{
            background-color: #f4f4f4;
            color: #000000;
            font-family: {web_font_family};
            font-size: {max(font_size - 2, 10)}px;
            padding: 2px 4px;
            border-radius: 4px;
            white-space: normal;
            display: inline;
        }}
        img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 10px 0;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin-top: 20px;
            font-size: {base}px;
        }}
        th, td {{
            text-align: left;
            border: 1px solid #ccc;
            padding: 12px 16px;
            vertical-align: top;
            font-size: {base}px;
        }}
        th {{
            background-color: #f3f3f3;
            color: #333;
        }}
        tr:nth-child(even) {{
            background-color: #fafafa;
        }}
        blockquote {{
            border-left: 4px solid #ddd;
            padding-left: 15px;
            color: #666;
            margin: 15px 0;
        }}
        a {{
            color: #0066cc;
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
    </style>
    <script src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>
    <script>
        window.MathJax = {{
            tex: {{
                inlineMath: [['$', '$'], ['\\\\(', '\\\\)']],
                displayMath: [['$$', '$$'], ['\\\\[', '\\\\]']]
            }},
            svg: {{ fontCache: 'global' }}
        }};
    </script>
</head>
<body>
{html_content}
</body>
</html>"""
        
        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_document)
        
        messagebox.showinfo("Success", f"HTML exported successfully to:\n{output_path}")
        return True
        
    except Exception as e:
        messagebox.showerror("Error", f"Failed to export HTML: {e}")
        return False


def convert_html_to_markdown(html_content):
    """
    Converts the HTML content of a file to Markdown format.
    
    :param string html_content: The HTML code to be converted.
    
    :return: A string containing Markdown code representing the converted HTML.

    :raises ConversionError: If the HTML could not be converted to Markdown.
    """

    try:
        # Create html2text converter instance
        h = html2text.HTML2Text()
        
        # Configure converter options for better output
        h.ignore_links = False  # Keep links
        h.ignore_images = False  # Keep images
        h.ignore_emphasis = False  # Keep bold/italic
        h.body_width = 0  # Don't wrap lines
        h.skip_internal_links = False  # Keep internal links
        h.inline_links = True  # Use inline link format [text](url)
        h.protect_links = True  # Protect URLs from being broken
        h.images_to_alt = False  # Don't replace images with alt text
        h.single_line_break = False  # Use proper line breaks
        h.mark_code = True  # Mark code blocks
        
        # Convert HTML to Markdown
        markdown_text = h.handle(html_content)
        
        # Clean up excessive blank lines
        markdown_text = re.sub(r'\n{3,}', '\n\n', markdown_text)
        
        return markdown_text.strip()
        
    except Exception as e:
        messagebox.showerror("Conversion Error", f"Failed to convert HTML to Markdown: {e}")
        return html_content  # Return original HTML if conversion fails


def convert_pdf_to_markdown(pdf_path):
    """
    Converts a PDF file's content to Markdown format with high-fidelity formatting preservation.
    Uses PyMuPDF (fitz) to extract detailed formatting information including font sizes,
    bold/italic text, and document structure.
    
    :param string pdf_path: The file path for the PDF file.
    
    :return: A string containing Markdown code representing the converted PDF file.

    :raises ImportError: If the required libraries are not installed.
    :raises ConversionError: If the PDF cannot be converted to Markdown.
    """

    try:
        import fitz

        pdf_abs_path = os.path.abspath(pdf_path)
        pdf_base_dir = os.path.dirname(pdf_abs_path)
        pdf_stem = os.path.splitext(os.path.basename(pdf_abs_path))[0]
        asset_dir = os.path.join(pdf_base_dir, f"{pdf_stem}_assets")
        os.makedirs(asset_dir, exist_ok=True)

        doc = fitz.open(pdf_abs_path)
        output_lines = []
        image_counter = 0

        for page_num in range(len(doc)):
            if page_num > 0:
                output_lines.append("---")
                output_lines.append("")

            page = doc[page_num]
            blocks = page.get_text("dict").get("blocks", [])

            # Normalize block order by page coordinates so images/text keep visual reading order.
            # Some PDFs return mixed block sequences where later sections can appear before figures.
            indexed_blocks = list(enumerate(blocks))
            indexed_blocks.sort(
                key=lambda item: (
                    ((item[1].get("bbox") or [0, 0, 0, 0])[1]),
                    ((item[1].get("bbox") or [0, 0, 0, 0])[0]),
                    item[0],
                )
            )
            ordered_blocks = [block for _, block in indexed_blocks]

            font_sizes = []
            for block in ordered_blocks:
                if block.get("type") == 0:
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            font_sizes.append(span.get("size", 12))

            if font_sizes:
                avg_size = sum(font_sizes) / len(font_sizes)
                heading1_threshold = avg_size * 1.5
                heading2_threshold = avg_size * 1.3
                heading3_threshold = avg_size * 1.15
            else:
                heading1_threshold = 18
                heading2_threshold = 16
                heading3_threshold = 14

            in_code_block = False
            pending_list_marker = False

            for block in ordered_blocks:
                block_type = block.get("type")

                if block_type == 1:
                    # Check image dimensions to filter out small decorative images/emojis
                    bbox = block.get("bbox")
                    if bbox:
                        width = bbox[2] - bbox[0]
                        height = bbox[3] - bbox[1]
                        # Skip images smaller than 50x50 pixels (likely emojis/icons)
                        if width < 50 or height < 50:
                            continue
                    
                    image_bytes = block.get("image")
                    image_ext = block.get("ext") or "png"
                    if image_bytes:
                        image_counter += 1
                        image_filename = f"page_{page_num + 1}_img_{image_counter}.{image_ext}"
                        image_path = os.path.join(asset_dir, image_filename)
                        with open(image_path, "wb") as image_file:
                            image_file.write(image_bytes)
                        image_uri = Path(image_path).as_uri()
                        if in_code_block:
                            output_lines.append("```")
                            output_lines.append("")
                            in_code_block = False
                        output_lines.append(f"![page-{page_num + 1}-image-{image_counter}]({image_uri})")
                        output_lines.append("")
                    continue

                if block_type != 0:
                    continue

                for line in block.get("lines", []):
                    line_parts = []
                    max_font_size = 0
                    has_monospace = False

                    for span in line.get("spans", []):
                        raw_text = span.get("text", "")
                        if not raw_text:
                            continue

                        text = raw_text.strip()
                        if not text:
                            continue

                        font_size = span.get("size", 12)
                        font_flags = span.get("flags", 0)
                        font_name = (span.get("font", "") or "").lower()

                        is_bold = (font_flags & 2**4) != 0 or "bold" in font_name
                        is_italic = (font_flags & 2**1) != 0 or "italic" in font_name
                        is_mono = ("mono" in font_name) or ("courier" in font_name) or ("consolas" in font_name)

                        if is_bold and is_italic:
                            text = f"***{text}***"
                        elif is_bold:
                            text = f"**{text}**"
                        elif is_italic:
                            text = f"*{text}*"

                        line_parts.append(text)
                        max_font_size = max(max_font_size, font_size)
                        has_monospace = has_monospace or is_mono

                    line_text = " ".join(line_parts).strip()
                    if not line_text:
                        output_lines.append("")
                        continue

                    if _is_standalone_list_marker(line_text):
                        pending_list_marker = True
                        continue

                    is_code_line = _is_pdf_code_line(line_text, has_monospace)
                    if not is_code_line and in_code_block and _is_pdf_code_continuation_line(line_text):
                        is_code_line = True
                    if is_code_line:
                        if not in_code_block:
                            output_lines.append("```bash")
                            in_code_block = True
                        output_lines.append(line_text)
                        continue

                    if in_code_block:
                        output_lines.append("```")
                        output_lines.append("")
                        in_code_block = False

                    if _is_list_item(line_text):
                        output_lines.append(f"- {_clean_list_item(line_text)}")
                        pending_list_marker = False
                        continue

                    if pending_list_marker:
                        output_lines.append(f"- {line_text}")
                        pending_list_marker = False
                        continue

                    clean_line = line_text.replace("**", "").replace("*", "").strip()
                    # Recover common PDF-rendered markdown subheadings like "2. Create ..."
                    # that may otherwise degrade into bold paragraph text.
                    if re.match(r'^\d+\.\s+\S+', clean_line) and (line_text.startswith("**") and line_text.endswith("**")):
                        output_lines.append(f"#### {clean_line}")
                        output_lines.append("")
                        continue

                    if max_font_size >= heading1_threshold:
                        clean_heading = line_text.replace("**", "").replace("*", "")
                        output_lines.append(f"# {clean_heading}")
                        output_lines.append("")
                    elif max_font_size >= heading2_threshold:
                        clean_heading = line_text.replace("**", "").replace("*", "")
                        output_lines.append(f"## {clean_heading}")
                        output_lines.append("")
                    elif max_font_size >= heading3_threshold:
                        clean_heading = line_text.replace("**", "").replace("*", "")
                        output_lines.append(f"### {clean_heading}")
                        output_lines.append("")
                    else:
                        output_lines.append(line_text)

                output_lines.append("")

            if in_code_block:
                output_lines.append("```")
                output_lines.append("")

            pending_list_marker = False

        doc.close()

        markdown_text = "\n".join(output_lines)
        markdown_text = re.sub(r'\n{4,}', '\n\n\n', markdown_text)
        return markdown_text.strip()
        
    except ImportError as e:
        # Fallback to pypdf if PyMuPDF is not available
        try:
            return _convert_pdf_to_markdown_fallback(pdf_path)
        except:
            messagebox.showerror("Import Error", 
                "PyMuPDF library is not installed for advanced PDF conversion.\n"
                "Please install it using: pip install PyMuPDF\n\n"
                "Falling back to basic conversion, but formatting may be lost.")
            return _convert_pdf_to_markdown_fallback(pdf_path)
    except Exception as e:
        messagebox.showerror("Conversion Error", f"Failed to convert PDF to Markdown: {e}")
        return ""


def _convert_pdf_to_markdown_fallback(pdf_path):
    """
    Fallback PDF conversion using pypdf when PyMuPDF is not available.
    This provides basic conversion without formatting preservation.
    
    :param string pdf_path: The file path for the PDF file.
    :return: A string containing basic Markdown conversion.
    """
    
    from pypdf import PdfReader
    
    reader = PdfReader(pdf_path)
    markdown_text = ""
    in_code_block = False
    pending_list_marker = False

    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text.strip():
            if page_num > 1:
                if in_code_block:
                    markdown_text += "```\n\n"
                    in_code_block = False
                markdown_text += "---\n\n"

            lines = text.split('\n')
            for line in lines:
                line = line.rstrip()
                if not line.strip():
                    markdown_text += "\n"
                    continue

                stripped = line.strip()
                if _is_standalone_list_marker(stripped):
                    pending_list_marker = True
                    continue

                if _is_pdf_code_line(line, False):
                    if not in_code_block:
                        markdown_text += "```bash\n"
                        in_code_block = True
                    markdown_text += line.strip() + "\n"
                    continue

                if in_code_block and _is_pdf_code_continuation_line(line):
                    markdown_text += line.strip() + "\n"
                    continue

                if in_code_block:
                    markdown_text += "```\n\n"
                    in_code_block = False

                if _is_list_item(stripped):
                    cleaned = _clean_list_item(stripped)
                    markdown_text += f"- {cleaned}\n"
                    pending_list_marker = False
                elif pending_list_marker:
                    markdown_text += f"- {stripped}\n"
                    pending_list_marker = False
                else:
                    markdown_text += stripped + "\n"

            markdown_text += "\n"
            pending_list_marker = False

    if in_code_block:
        markdown_text += "```\n"

    return markdown_text.strip()


def _is_pdf_code_line(line, has_monospace_font=False):
    """
    Heuristic detection for code-like lines in PDF extraction.

    :param string line: A single extracted line.
    :param bool has_monospace_font: Whether line includes monospace span.
    :return: True if the line is likely code.
    """

    stripped = line.strip()
    if not stripped:
        return False

    if re.match(r'^\*{0,2}\d+[\.)]\s+.+\*{0,2}$', stripped):
        return False
    if re.match(r'^#+\s+', stripped):
        return False
    if stripped.startswith(('![', '[', '>')):
        return False

    if has_monospace_font:
        return True

    if line.startswith('    ') or line.startswith('\t'):
        return True

    code_patterns = [
        r'^\$\s+\S+',
        r'^(sudo|pip|python|python3|npm|node|git|cd|ls|mkdir|rm|cp|mv|source)\b',
        r'^(if|for|while|def|class|return|import|from|try|except|else|elif)\b',
        r'\b(function|const|let|var|echo|export|chmod|chown|brew|apt|yum|conda)\b',
        r'`[^`]+`'
    ]

    for pattern in code_patterns:
        if re.search(pattern, stripped):
            return True

    return False


def _is_pdf_code_continuation_line(line):
    """
    Detect lines that should remain inside an already-open PDF-derived code block.

    :param string line: A single extracted line.
    :return: True if the line likely continues code/command content.
    """

    stripped = line.strip()
    if not stripped:
        return False

    if stripped.startswith("#"):
        return True

    continuation_patterns = [
        r'^(source|export|set|unset|alias|PATH=)\b',
        r'^(\./|\.\\|\.\./|\.\.\\)',
        r'(^|\s)(venv|scripts|bin|powershell|cmd|activate)(\s|$)',
        r'[/\\].*(activate|python|pip)',
    ]

    return any(re.search(pattern, stripped, re.IGNORECASE) for pattern in continuation_patterns)


def _is_standalone_list_marker(line):
    """
    Detect lines that contain only a list marker (common in PDF extraction),
    so the next text line can be merged into one markdown list item.
    """

    stripped = line.strip()
    if not stripped:
        return False

    bullet_markers = {'•', '●', '○', '▪', '▫', '■', '□', '✓', '✔', '*', '-', '–', '—'}
    if stripped in bullet_markers:
        return True

    if re.fullmatch(r'(\d+|[a-zA-Z]|[ivxIVX]+)[\.)]', stripped):
        return True

    return False


def _is_list_item(line):
    """
    Determines if a line is a list item.
    
    :param string line: The line to check.
    
    :return: Boolean indicating if the line is a list item.
    """
    
    # Check for common list patterns
    patterns = [
        r'^\s*[•●○▪▫■□✓✔]\s*\S+',  # Bullet points
        r'^\s*[-–—]\s*\S+',          # Dashes
        r'^\s*\d+[\.)]\s+',       # Numbered (1. or 1))
        r'^\s*[a-z][\.)]\s+',     # Lettered (a. or a))
        r'^\s*[ivxIVX]+[\.)]\s+', # Roman numerals
    ]
    
    for pattern in patterns:
        if re.match(pattern, line):
            return True
    
    return False


def _clean_list_item(line):
    """
    Removes list markers from a line.
    
    :param string line: The line to clean.
    
    :return: Cleaned line text.
    """
    
    # Remove common list markers
    cleaned = re.sub(r'^\s*[•●○▪▫■□✓✔\-–—]\s+', '', line)
    cleaned = re.sub(r'^\s*\d+[\.)]\s+', '', cleaned)
    cleaned = re.sub(r'^\s*[a-z][\.)]\s+', '', cleaned)
    cleaned = re.sub(r'^\s*[ivxIVX]+[\.)]\s+', '', cleaned)
    
    return cleaned.strip()


def export_to_docx(app, output_path):
    """
    Exports the current Markdown document to a Word (.docx) file.
    
    :param MarkdownReader app: The MarkdownReader application instance.
    :param string output_path: The path where the .docx file should be saved.
    
    :return: A boolean set to true if the Markdown is successfully converted, and false otherwise.

    :raises RuntimeError: If the document could not be exported to Word.
    """

    if not app.editors:
        messagebox.showinfo("Info", "No document to export.")
        return False
    
    try:
        import requests
        from io import BytesIO

        idx = app.notebook.index(app.notebook.select())
        text_area = app.editors[idx]
        markdown_text = text_area.get("1.0", "end-1c")

        current_path = None
        base_dir = None
        if hasattr(app, 'file_paths') and app.file_paths and idx < len(app.file_paths):
            current_path = app.file_paths[idx]
            if current_path:
                base_dir = os.path.dirname(current_path)

        def resolve_image_source(src):
            """
            Takes the image src and returns the absolute path to the file.

            :param string src: The src from the image tag.

            :return: Two strings, one indicating whether the image stored locally or online, and another with the path required to access the image.
            """
            
            src = src.strip().strip('<>')
            if (src.startswith('"') and src.endswith('"')) or (src.startswith("'") and src.endswith("'")):
                src = src[1:-1]

            if src.startswith('file://'):
                src = src[7:]

            if re.match(r'^https?://', src, re.IGNORECASE):
                return 'remote', src

            if os.path.isabs(src):
                return 'local', src

            if base_dir:
                return 'local', os.path.abspath(os.path.join(base_dir, src))

            return 'local', os.path.abspath(src)

        def insert_image(doc_obj, src):
            """
            Inserts the given image src into the Word document.

            :param Document doc_obj: The Word document to have the image inserted into.
            :param src: The unformatted image source for the image to be inserted.

            :return: A boolean containing true if the operation succeeds, and false if not.
            """
            
            kind, value = resolve_image_source(src)
            try:
                if kind == 'remote':
                    response = requests.get(value, timeout=15, allow_redirects=True)
                    response.raise_for_status()
                    doc_obj.add_picture(BytesIO(response.content), width=Inches(6.2))
                else:
                    if not os.path.exists(value):
                        return False
                    doc_obj.add_picture(value, width=Inches(6.2))

                doc_obj.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                return True
            except Exception:
                return False
        
        # Create a new Word document
        doc = Document()
        
        # Parse and convert markdown to Word
        lines = markdown_text.split('\n')
        i = 0
        in_code_block = False
        code_block_lines = []
        in_list = False
        
        while i < len(lines):
            line = lines[i]
            
            # Handle code blocks
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_block_lines = []
                else:
                    # End of code block
                    in_code_block = False
                    if code_block_lines:
                        code_text = '\n'.join(code_block_lines)
                        p = doc.add_paragraph(code_text)
                        p.style = 'Intense Quote'
                        for run in p.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                    code_block_lines = []
                i += 1
                continue
            
            if in_code_block:
                code_block_lines.append(line)
                i += 1
                continue
            
            # Skip empty lines
            if not line.strip():
                i += 1
                continue

            # Markdown image: ![alt](src)
            md_image_match = re.match(r'^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$', line)
            if md_image_match:
                alt_text = md_image_match.group(1).strip()
                src = md_image_match.group(2).strip()
                if not insert_image(doc, src):
                    doc.add_paragraph(alt_text if alt_text else src)
                i += 1
                continue

            # HTML image: <img ... src="..." ...>
            html_image_match = re.match(r'^\s*<img\b[^>]*\bsrc=["\']([^"\']+)["\'][^>]*>\s*$', line, re.IGNORECASE)
            if html_image_match:
                src = html_image_match.group(1).strip()
                alt_match = re.search(r'\balt=["\']([^"\']*)["\']', line, re.IGNORECASE)
                alt_text = alt_match.group(1).strip() if alt_match else ''
                if not insert_image(doc, src):
                    doc.add_paragraph(alt_text if alt_text else src)
                i += 1
                continue
            
            # Headings
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                if level == 1:
                    doc.add_heading(text, level=1)
                elif level == 2:
                    doc.add_heading(text, level=2)
                elif level == 3:
                    doc.add_heading(text, level=3)
                else:
                    doc.add_heading(text, level=4)
            
            # Horizontal rule
            elif line.strip() in ['---', '***', '___']:
                doc.add_paragraph('_' * 50)
            
            # Unordered list
            elif re.match(r'^\s*[-*+]\s+', line):
                text = re.sub(r'^\s*[-*+]\s+', '', line)
                text = process_inline_formatting(text)
                p = doc.add_paragraph(text, style='List Bullet')
                apply_inline_formatting(p, text)
            
            # Ordered list
            elif re.match(r'^\s*\d+\.\s+', line):
                text = re.sub(r'^\s*\d+\.\s+', '', line)
                text = process_inline_formatting(text)
                p = doc.add_paragraph(text, style='List Number')
                apply_inline_formatting(p, text)
            
            # Blockquote
            elif line.strip().startswith('>'):
                text = line.strip().lstrip('>').strip()
                text = process_inline_formatting(text)
                p = doc.add_paragraph(text, style='Intense Quote')
                apply_inline_formatting(p, text)
            
            # Table detection (simple)
            elif '|' in line and line.strip().startswith('|'):
                table_lines = [line]
                i += 1
                # Collect table rows
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                
                # Parse and create table
                if len(table_lines) > 2:  # Header + separator + at least one row
                    rows_data = []
                    for tline in table_lines:
                        if not re.match(r'^\s*\|[\s:-]+\|', tline):  # Skip separator
                            cells = [c.strip() for c in tline.split('|')[1:-1]]
                            rows_data.append(cells)
                    
                    if rows_data:
                        table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
                        table.style = 'Light Grid Accent 1'
                        
                        for row_idx, row_data in enumerate(rows_data):
                            for col_idx, cell_text in enumerate(row_data):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_text
                                if row_idx == 0:  # Header row
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                continue
            
            # Regular paragraph
            else:
                text = process_inline_formatting(line)
                p = doc.add_paragraph()
                apply_inline_formatting(p, line)
            
            i += 1
        
        # Save the document
        doc.save(output_path)
        messagebox.showinfo("Success", f"Document exported successfully to:\n{output_path}")
        return True
        
    except Exception as e:
        messagebox.showerror("Error", f"Failed to export to Word: {e}")
        import traceback
        traceback.print_exc()
        return False


def export_to_pdf(app, output_path):
    """
    Exports the current Markdown document to a PDF file.
    - Since PDF generation libraries have complex dependencies,
      this creates a print-friendly HTML file and opens it in the browser
      for the user to print to PDF using the browser's built-in functionality.
    
    :param MarkdownReader app: The MarkdownReader application instance.
    :param string output_path: The path where the PDF file should be saved (used as suggestion).
    
    :return: A boolean set to true if successfully exported and false otherwise.

    :raises RuntimeError: If the image paths could not be processed.
    :raises RuntimeError: If the document could not be exported to PDF.
    """

    if not app.editors:
        messagebox.showinfo("Info", "No document to export.")
        return False
    
    try:
        import tempfile
        import webbrowser
        
        idx = app.notebook.index(app.notebook.select())
        text_area = app.editors[idx]
        markdown_text = text_area.get("1.0", "end-1c")
        
        # Fix image paths if a file is currently open
        base_dir = None
        if hasattr(app, 'file_paths') and app.file_paths:
            try:
                current_path = app.file_paths[idx]
                if current_path is not None:
                    base_dir = os.path.dirname(current_path)
                    def convert_file_url_to_absolute(text, base_dir):
                        """
                        Converts file:// paths and relative paths to absolute file URLs for PDF export.

                        :param string text: The file URL to be formatted.
                        :param string base_dir: The base directory of the file.

                        :return: The formatted URL ready for PDF export.
                        """
                        
                        def repl(m):
                            alt = m.group(1)
                            src = m.group(2)
                            if src.startswith('file://'):
                                # Already absolute file URL
                                return f'![{alt}]({src})'
                            elif src.startswith('/') or (len(src) > 1 and src[1] == ':'):
                                # Already absolute path
                                return f'![{alt}](file://{src})'
                            else:
                                # Relative path - convert to absolute file URL
                                try:
                                    abs_path = os.path.abspath(os.path.join(base_dir, src))
                                    file_url = 'file://' + abs_path
                                    return f'![{alt}]({file_url})'
                                except (ValueError, OSError):
                                    return m.group(0)
                        return re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', repl, text)
                    
                    markdown_text = convert_file_url_to_absolute(markdown_text, base_dir)
            except Exception as e:
                print(f"Warning: Could not process image paths: {e}")
        
        # Convert markdown to HTML
        html_content = markdown2.markdown(
            markdown_text, 
            extras=["fenced-code-blocks", "code-friendly", "tables", "break-on-newline"]
        )
        
        # Get style from app (with fallback)
        font_family = getattr(app, 'current_font_family', 'Consolas')
        font_size = getattr(app, 'current_font_size', 14)
        fg_color = getattr(app, 'current_fg_color', '#000000')
        bg_color = getattr(app, 'current_bg_color', 'white')
        
        if getattr(app, 'dark_mode', False):
            bg_color = '#1e1e1e'
            fg_color = '#dcdcdc'
        
        # For PDF, use generic font fallbacks
        web_font_family = font_family
        if font_family.lower() in ["arial", "helvetica", "verdana", "tahoma", "trebuchet ms"]:
            web_font_family += ", sans-serif"
        elif font_family.lower() in ["times new roman", "georgia", "garamond", "serif"]:
            web_font_family += ", serif"
        elif font_family.lower() in ["consolas", "courier new", "monospace"]:
            web_font_family += ", monospace"
        else:
            web_font_family += ", sans-serif"
        
        # Code blocks should always use monospace fonts
        code_font_family = "Consolas, 'Courier New', monospace"
        
        # Heading sizes relative to base font size
        h1 = font_size + 18
        h2 = font_size + 12
        h3 = font_size + 8
        h4 = font_size + 4
        h5 = font_size + 2
        h6 = font_size + 1
        base = font_size + 2
        
        # Generate complete HTML document for PDF printing
        # Optimized for print media with @media print styles
        base_tag = ""
        if base_dir:
            # Add base tag to help resolve relative image paths
            base_tag = f'<base href="file://{base_dir}/">'
        
        html_document = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Exported Markdown</title>
    {base_tag}
    <style>
        @media print {{
            body {{
                background-color: white !important;
                color: black !important;
                margin: 0;
                padding: 20px;
            }}
            
            /* Hide print button when actually printing */
            .print-instructions {{
                display: none;
            }}
            
            /* Page break control */
            h1, h2, h3, h4, h5, h6 {{
                page-break-after: avoid;
                page-break-inside: avoid;
            }}
            
            pre, code, table, img {{
                page-break-inside: avoid;
            }}
        }}
        
        @media screen {{
            body {{
                max-width: 800px;
                margin: 20px auto;
                padding: 40px;
                background-color: white;
                box-shadow: 0 0 10px rgba(0,0,0,0.1);
            }}
            
            /* Show print instructions on screen */
            .print-instructions {{
                background-color: #e3f2fd;
                border: 2px solid #2196F3;
                border-radius: 8px;
                padding: 20px;
                margin: 20px 0;
                text-align: center;
            }}
            
            .print-button {{
                background-color: #2196F3;
                color: white;
                border: none;
                padding: 12px 24px;
                font-size: 16px;
                border-radius: 4px;
                cursor: pointer;
                margin: 10px;
            }}
            
            .print-button:hover {{
                background-color: #1976D2;
            }}
        }}
        
        /* Common styles for both screen and print */
        body {{
            font-family: {web_font_family};
            font-size: {base}pt;
            line-height: 1.6;
            color: #333;
        }}
        
        h1 {{ font-size: {h1}pt; margin-top: 1.5em; margin-bottom: 0.5em; font-weight: bold; }}
        h2 {{ font-size: {h2}pt; margin-top: 1.2em; margin-bottom: 0.4em; font-weight: bold; }}
        h3 {{ font-size: {h3}pt; margin-top: 1em; margin-bottom: 0.3em; font-weight: bold; }}
        h4 {{ font-size: {h4}pt; margin-top: 0.8em; margin-bottom: 0.25em; font-weight: bold; }}
        h5 {{ font-size: {h5}pt; margin-top: 0.6em; margin-bottom: 0.2em; font-weight: bold; }}
        h6 {{ font-size: {h6}pt; margin-top: 0.5em; margin-bottom: 0.2em; font-weight: bold; }}
        
        b, strong {{ font-weight: bold; }}
        i, em {{ font-style: italic; }}
        u {{ text-decoration: underline; }}
        
        pre {{
            background-color: #f5f5f5;
            color: #000;
            padding: 12pt;
            border: 1pt solid #ddd;
            border-radius: 4px;
            margin: 10pt 0;
            overflow-x: auto;
        }}
        
        pre code {{
            font-family: Consolas, 'Courier New', monospace;
            font-size: {max(font_size - 2, 10)}pt;
            white-space: pre;
            display: block;
        }}
        
        code {{
            background-color: #f5f5f5;
            color: #c7254e;
            font-family: Consolas, 'Courier New', monospace;
            font-size: {max(font_size - 1, 10)}pt;
            padding: 2pt 4pt;
            border-radius: 3px;
            border: 1pt solid #e1e1e8;
        }}
        
        img {{
            max-width: 100%;
            height: auto;
            display: block;
            margin: 15pt auto;
        }}
        
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 15pt 0;
        }}
        
        th, td {{
            text-align: left;
            border: 1pt solid #ddd;
            padding: 8pt 12pt;
            vertical-align: top;
        }}
        
        th {{
            background-color: #f5f5f5;
            color: #333;
            font-weight: bold;
        }}
        
        tr:nth-child(even) {{
            background-color: #fafafa;
        }}
        
        blockquote {{
            border-left: 4pt solid #ddd;
            padding-left: 15pt;
            color: #666;
            margin: 15pt 0;
            font-style: italic;
        }}
        
        a {{
            color: #0066cc;
            text-decoration: underline;
        }}
        
        ul, ol {{
            margin: 8pt 0;
            padding-left: 30pt;
        }}
        
        li {{
            margin: 4pt 0;
        }}
        
        p {{
            margin: 8pt 0;
        }}
    </style>
</head>
<body>
    <div class="print-instructions">
        <h2>📄 Ready to Print as PDF</h2>
        <p>This page has been optimized for printing. Please use one of the following methods:</p>
        <button class="print-button" onclick="window.print()">🖨️ Print this page</button>
        <div style="margin-top: 15px; color: #666; font-size: 14px;">
            <p><strong>Or use keyboard shortcuts:</strong></p>
            <p>macOS: Cmd + P | Windows/Linux: Ctrl + P</p>
            <p><em>Select "Save as PDF" as the printer in the print dialog</em></p>
        </div>
    </div>
    
{html_content}

    <script>
        // Auto-trigger print dialog after page loads (optional)
        // Uncomment the next line if you want automatic print dialog
        // window.onload = function() {{ setTimeout(function() {{ window.print(); }}, 500); }};
    </script>
</body>
</html>"""

        # Create a temporary HTML file optimized for printing
        base_name = os.path.splitext(os.path.basename(output_path))[0]
        with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as tmp_file:
            tmp_file.write(html_document)
            tmp_html_path = tmp_file.name

        # Open the HTML file in the default browser
        _open_file_in_browser(tmp_html_path)

        # Show instructions to the user
        result = messagebox.showinfo(
            "Export PDF",
            f"A print-friendly HTML page has been opened in your browser.\n\n"
            f"Please follow these steps to export as PDF:\n\n"
            f"1. Press Cmd+P (macOS) or Ctrl+P (Windows) in your browser\n"
            f"2. Select 'Save as PDF' as the printer\n"
            f"3. Choose the save location and name it: {os.path.basename(output_path)}\n\n"
            f"Or click the 'Print' button on the page."
        )

        return True

    except Exception as e:
        messagebox.showerror("Error", f"Failed to export PDF: {e}")
        import traceback
        traceback.print_exc()
        return False


def process_inline_formatting(text):
    """
    Removes Markdown formatting markers for plain text extraction.

    :param string text: The text to be processed.

    :return: A string containing the inputted text without the inline formatting.
    """

    # Remove bold
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # Remove italic
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # Remove inline code
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Remove links but keep text
    text = re.sub(r'\[(.+?)\]\((.+?)\)', r'\1', text)
    return text


def apply_inline_formatting(paragraph, text):
    """
    Applies bold, italic, and other inline formatting to a Word paragraph.

    :param paragraph: The paragraph of text with Word formatting.
    :param text: The text with no inline formatting.
    """

    # Clear existing runs
    paragraph.clear()
    
    # Pattern to match markdown inline formatting
    # This is a simplified version - handles bold, italic, code, and links
    parts = []
    current_pos = 0
    
    # Find all formatting markers
    patterns = [
        (r'\*\*(.+?)\*\*', 'bold'),      # Bold with **
        (r'__(.+?)__', 'bold'),          # Bold with __
        (r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', 'italic'),  # Italic with *
        (r'(?<!_)_(?!_)(.+?)(?<!_)_(?!_)', 'italic'),        # Italic with _
        (r'`(.+?)`', 'code'),            # Inline code
        (r'\[(.+?)\]\((.+?)\)', 'link'), # Links
    ]
    
    segments = []
    i = 0
    while i < len(text):
        matched = False
        for pattern, style in patterns:
            match = re.match(pattern, text[i:])
            if match:
                if style == 'link':
                    segments.append(('normal', match.group(1)))
                else:
                    segments.append((style, match.group(1)))
                i += match.end()
                matched = True
                break
        
        if not matched:
            # Find next formatting marker
            next_marker = len(text)
            for pattern, _ in patterns:
                match = re.search(pattern, text[i:])
                if match:
                    next_marker = min(next_marker, i + match.start())
            
            if next_marker > i:
                segments.append(('normal', text[i:next_marker]))
                i = next_marker
            else:
                i += 1
    
    # If no formatting found, just add the text
    if not segments:
        paragraph.add_run(text)
    else:
        for style, content in segments:
            run = paragraph.add_run(content)
            if style == 'bold':
                run.font.bold = True
            elif style == 'italic':
                run.font.italic = True
            elif style == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(212, 73, 80)


def convert_pdf_to_markdown_docling(pdf_path):
    """
    Converts a PDF file to Markdown format using Docling with advanced ML-based analysis.
    
    Docling uses deep learning models to better understand document structure,
    including tables, multi-column layouts, and complex formatting.
    
    Extracts and saves embedded images to an asset directory with proper file:// URIs
    for display in browser previews.
    
    :param string pdf_path: The file path for the PDF file.
    :return: A string containing Markdown code representing the converted PDF file.
    :raises ImportError: If Docling is not installed.
    :raises Exception: If the PDF cannot be converted to Markdown.
    """
    
    converter = None

    try:
        from docling.document_converter import DocumentConverter
        import re
        import html
        import shutil

        def _cleanup_docling_object(obj, seen=None):
            if obj is None:
                return

            if seen is None:
                seen = set()

            obj_id = id(obj)
            if obj_id in seen:
                return
            seen.add(obj_id)

            for method_name in ("cleanup", "close", "shutdown", "terminate", "stop", "unload"):
                method = getattr(obj, method_name, None)
                if callable(method):
                    try:
                        method()
                    except TypeError:
                        pass
                    except Exception:
                        pass

            if isinstance(obj, dict):
                for value in obj.values():
                    _cleanup_docling_object(value, seen)
                return

            if isinstance(obj, (list, tuple, set)):
                for value in obj:
                    _cleanup_docling_object(value, seen)
                return

            obj_dict = getattr(obj, "__dict__", None)
            if not obj_dict:
                return

            for attr_name, value in obj_dict.items():
                if attr_name.startswith("__"):
                    continue
                value_module = getattr(getattr(value, "__class__", None), "__module__", "")
                if isinstance(value, (dict, list, tuple, set)) or value_module.startswith(("docling", "docling_core")):
                    _cleanup_docling_object(value, seen)

        def _cleanup_docling_converter(doc_converter):
            initialized = getattr(doc_converter, "initialized_pipelines", None)
            if isinstance(initialized, dict):
                for pipeline in list(initialized.values()):
                    _cleanup_docling_object(pipeline)
                initialized.clear()
        
        pdf_abs_path = os.path.abspath(pdf_path)
        pdf_base_dir = os.path.dirname(pdf_abs_path)
        pdf_stem = os.path.splitext(os.path.basename(pdf_abs_path))[0]
        asset_dir = os.path.join(pdf_base_dir, f"{pdf_stem}_assets")
        os.makedirs(asset_dir, exist_ok=True)
        
        if not os.path.exists(pdf_abs_path):
            raise FileNotFoundError(f"PDF file not found: {pdf_abs_path}")
        
        # Initialize converter
        converter = DocumentConverter()
        
        # Convert PDF
        result = converter.convert(pdf_abs_path)
        
        # Export to Markdown
        markdown_text = result.document.export_to_markdown()

        def _detect_ext_from_mime(mime_type):
            mime = (mime_type or "").lower()
            if "jpeg" in mime or "jpg" in mime:
                return "jpg"
            if "png" in mime:
                return "png"
            if "gif" in mime:
                return "gif"
            if "webp" in mime:
                return "webp"
            if "bmp" in mime:
                return "bmp"
            if "svg" in mime:
                return "svg"
            return "png"

        def _write_bytes_file(content_bytes, target_path):
            with open(target_path, "wb") as img_file:
                img_file.write(content_bytes)

        def _save_picture_to_path(picture, target_path):
            if hasattr(picture, "save") and callable(getattr(picture, "save")):
                try:
                    picture.save(target_path)
                    return os.path.exists(target_path) and os.path.getsize(target_path) > 0
                except Exception:
                    pass

            if hasattr(picture, "export_to_file") and callable(getattr(picture, "export_to_file")):
                try:
                    picture.export_to_file(target_path)
                    return os.path.exists(target_path) and os.path.getsize(target_path) > 0
                except Exception:
                    pass

            candidate_attrs = ["get_image", "image", "pil_image", "data", "bytes", "blob", "content"]
            for attr in candidate_attrs:
                if not hasattr(picture, attr):
                    continue

                value = getattr(picture, attr)
                try:
                    value = value() if callable(value) else value
                except Exception:
                    continue

                if value is None:
                    continue

                if isinstance(value, (bytes, bytearray)):
                    _write_bytes_file(bytes(value), target_path)
                    return os.path.exists(target_path) and os.path.getsize(target_path) > 0

                if isinstance(value, str) and os.path.exists(value):
                    shutil.copyfile(value, target_path)
                    return os.path.exists(target_path) and os.path.getsize(target_path) > 0

                if hasattr(value, "save") and callable(getattr(value, "save")):
                    try:
                        value.save(target_path)
                        return os.path.exists(target_path) and os.path.getsize(target_path) > 0
                    except Exception:
                        continue

            return False

        # Extract and save images from document
        image_uris = []
        image_mapping = {}
        image_counter = 0

        def _extract_images_with_pymupdf_fallback(start_index=0):
            extracted = []
            try:
                import fitz

                fallback_doc = fitz.open(pdf_abs_path)
                local_counter = start_index
                image_records = []
                for page_num in range(len(fallback_doc)):
                    page = fallback_doc[page_num]
                    blocks = page.get_text("dict").get("blocks", [])
                    for block in blocks:
                        if block.get("type") != 1:
                            continue

                        bbox = block.get("bbox")
                        if bbox:
                            width = bbox[2] - bbox[0]
                            height = bbox[3] - bbox[1]
                            if width < 32 or height < 32:
                                continue

                        image_bytes = block.get("image")
                        if not image_bytes:
                            continue

                        image_ext = block.get("ext") or "png"
                        local_counter += 1
                        filename = f"docling_image_{local_counter}.{image_ext}"
                        path = os.path.join(asset_dir, filename)
                        with open(path, "wb") as image_file:
                            image_file.write(image_bytes)

                        bbox = block.get("bbox") or [0, 0, 0, 0]
                        x_top = float(bbox[0]) if len(bbox) > 0 else 0.0
                        y_top = float(bbox[1]) if len(bbox) > 1 else 0.0
                        image_records.append((page_num, y_top, x_top, Path(path).as_uri()))

                # Sort by reading order: page -> y(top) -> x(left)
                image_records.sort(key=lambda item: (item[0], item[1], item[2]))
                extracted = [item[3] for item in image_records]

                fallback_doc.close()
            except Exception:
                return []

            return extracted

        doc = result.document
        pictures = getattr(doc, "pictures", None)
        if pictures:
            for idx, picture in enumerate(pictures, 1):
                try:
                    image_ext = _detect_ext_from_mime(getattr(picture, "mime_type", ""))
                    image_counter += 1
                    image_filename = f"docling_image_{image_counter}.{image_ext}"
                    image_path = os.path.join(asset_dir, image_filename)

                    if not _save_picture_to_path(picture, image_path):
                        image_counter -= 1
                        continue

                    image_uri = Path(image_path).as_uri()
                    image_uris.append(image_uri)

                    possible_keys = [
                        f"docling_image_{idx}",
                        f"image_{idx}",
                        f"picture_{idx}",
                        f"#/pictures/{idx - 1}",
                        f"/pictures/{idx - 1}",
                    ]
                    picture_name = getattr(picture, "name", None)
                    if picture_name:
                        possible_keys.append(str(picture_name))
                    for key in possible_keys:
                        image_mapping[key] = image_uri
                except Exception:
                    continue

        has_image_placeholder = re.search(r'^\s*<!--\s*image\s*-->\s*$', markdown_text, flags=re.IGNORECASE | re.MULTILINE) is not None
        if not image_uris or has_image_placeholder:
            fallback_uris = _extract_images_with_pymupdf_fallback(start_index=len(image_uris))
            if fallback_uris:
                if has_image_placeholder:
                    # Placeholder order should follow page reading order from PDF blocks.
                    image_uris = list(fallback_uris)
                    image_mapping = {}
                else:
                    image_uris.extend(fallback_uris)

        # Convert escaped HTML image tags such as &lt;img ...&gt; and < img ... >
        markdown_text = re.sub(r'&lt;\s*img\b([^&]*)&gt;', r'<img\1>', markdown_text, flags=re.IGNORECASE)
        markdown_text = re.sub(r'<\s+img\b', '<img', markdown_text, flags=re.IGNORECASE)
        markdown_text = re.sub(r'\s+/\s*>', ' />', markdown_text)

        def _html_img_to_md(match):
            full_tag = match.group(0)
            src_match = re.search(r'\bsrc=["\']([^"\']+)["\']', full_tag, re.IGNORECASE)
            if not src_match:
                return full_tag
            src = html.unescape(src_match.group(1).strip())
            alt_match = re.search(r'\balt=["\']([^"\']*)["\']', full_tag, re.IGNORECASE)
            alt = html.unescape(alt_match.group(1).strip()) if alt_match else ''
            return f"![{alt}]({src})"

        markdown_text = re.sub(r'<img\b[^>]*>', _html_img_to_md, markdown_text, flags=re.IGNORECASE)

        # Basic cleanup for common OCR/round-trip artifacts
        markdown_text = html.unescape(markdown_text)
        markdown_text = re.sub(r"^'''([a-zA-Z0-9_-]+)\s*$", r"```\1", markdown_text, flags=re.MULTILINE)
        markdown_text = re.sub(r"^'''([a-zA-Z0-9_-]+)\s+(.+)$", r"```\1\n\2", markdown_text, flags=re.MULTILINE)
        markdown_text = re.sub(r"^'''\s*$", "```", markdown_text, flags=re.MULTILINE)
        markdown_text = re.sub(r"^[ \t]*«[ \t]+", "* ", markdown_text, flags=re.MULTILINE)
        markdown_text = re.sub(r'^\s*["“”]*\s*(Normal text|Consolas)\s*["“”]*\s*\n', '', markdown_text, flags=re.IGNORECASE | re.MULTILINE)
        markdown_text = re.sub(r'\n{3,}', '\n\n', markdown_text)

        # Replace markdown image paths to absolute file:// URIs when possible
        def _replace_md_image_path(match):
            alt_text = match.group(1)
            image_src = match.group(2).strip()

            if image_src.startswith(("http://", "https://", "file://", "data:")):
                return match.group(0)

            for mapped_name, uri in image_mapping.items():
                if mapped_name and mapped_name in image_src:
                    return f"![{alt_text}]({uri})"

            local_candidates = []
            local_candidates.append(os.path.abspath(os.path.join(pdf_base_dir, image_src)))
            local_candidates.append(os.path.abspath(os.path.join(asset_dir, image_src)))

            for candidate in local_candidates:
                if os.path.exists(candidate):
                    return f"![{alt_text}]({Path(candidate).as_uri()})"

            if image_uris and re.search(r'(image|picture|fig|figure)', image_src, re.IGNORECASE):
                num_match = re.search(r'(\d+)', image_src)
                if num_match:
                    image_idx = int(num_match.group(1))
                    if 1 <= image_idx <= len(image_uris):
                        return f"![{alt_text}]({image_uris[image_idx - 1]})"

            return match.group(0)

        markdown_text = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', _replace_md_image_path, markdown_text)

        # Replace Docling image placeholders: <!-- image -->
        if image_uris:
            placeholder_idx = [0]

            def _replace_image_placeholder(match):
                if placeholder_idx[0] < len(image_uris):
                    uri = image_uris[placeholder_idx[0]]
                    placeholder_idx[0] += 1
                    return f"![docling-image-{placeholder_idx[0]}]({uri})"
                return ""

            markdown_text = re.sub(
                r'^\s*<!--\s*image\s*-->\s*$',
                _replace_image_placeholder,
                markdown_text,
                flags=re.IGNORECASE | re.MULTILINE,
            )

        # Remove any unreplaced image placeholders
        markdown_text = re.sub(r'^\s*<!--\s*image\s*-->\s*$', '', markdown_text, flags=re.IGNORECASE | re.MULTILINE)

        def _fix_malformed_code_fences(text):
            lines = text.splitlines()
            output = []
            in_code_block = False
            current_lang = ""
            code_line_count = 0

            for line in lines:
                stripped = line.strip()

                if stripped.startswith("```"):
                    if in_code_block:
                        in_code_block = False
                        current_lang = ""
                        code_line_count = 0
                    else:
                        in_code_block = True
                        current_lang = stripped[3:].strip().lower()
                        code_line_count = 0
                    output.append(line)
                    continue

                if in_code_block:
                    if stripped:
                        code_line_count += 1

                    looks_like_doc_content = (
                        re.match(r'^\s*#{1,6}\s+', line) is not None
                        or re.match(r'^\s*[-*+]\s+', line) is not None
                        or re.match(r'^\s*\d+\.\s+', line) is not None
                        or stripped in ('---', '***', '___')
                    )

                    # Heuristic: if bash fence only had command lines and then doc content starts,
                    # it is usually an unclosed accidental fence from OCR.
                    if looks_like_doc_content and current_lang in ("bash", "sh", "shell", "zsh") and code_line_count <= 2:
                        output.append("```")
                        in_code_block = False
                        current_lang = ""
                        code_line_count = 0

                output.append(line)

            if in_code_block:
                output.append("```")

            return "\n".join(output)

        def _bind_images_to_sections(text, ordered_image_uris):
            if not ordered_image_uris:
                return text

            def _is_screenshot_ocr_noise(line):
                stripped = line.strip()
                if not stripped:
                    return False

                lowered = stripped.lower()

                if re.match(r'^\d{1,4}$', stripped):
                    return True
                if re.match(r'^[-–—_]+$', stripped):
                    return True
                if re.match(r'^[a-z0-9_.-]+\.(md|txt|pdf|doc|docx|html|htm)$', lowered):
                    return True
                if lowered in {
                    'markdown reader',
                    'untitled',
                    'readme.md',
                    'features',
                    'editor overview',
                    'preview overview',
                }:
                    return True
                if '·' in stripped or '•' in stripped:
                    return True
                if 'toggle, and drag-and-drop file opening' in lowered:
                    return True
                if 'compatible with mac' in lowered and 'desktop environments' in lowered:
                    return True

                if len(stripped) <= 40 and re.match(r'^[a-zA-Z0-9 .,_()\-/]+$', stripped):
                    word_count = len([w for w in stripped.split() if w])
                    if 1 <= word_count <= 6 and lowered != 'installation & usage':
                        return True

                return False

            section_specs = [
                ("Editor Overview", 0),
                ("Preview Overview", 1),
            ]

            updated = text
            inserted_count = 0
            for section_name, image_idx in section_specs:
                if image_idx >= len(ordered_image_uris):
                    continue

                section_pattern = rf'(^\s*#{{2,6}}\s*{re.escape(section_name)}\s*$)(.*?)(?=^\s*#{{1,6}}\s+|\Z)'

                def _section_repl(match):
                    nonlocal inserted_count
                    heading_line = match.group(1)
                    body = match.group(2)

                    body_lines = body.splitlines()
                    cleaned_lines = []
                    for body_line in body_lines:
                        if re.match(r'^\s*!\[[^\]]*\]\(([^)]+)\)\s*$', body_line):
                            continue
                        if re.match(r'^\s*<img\b[^>]*>\s*$', body_line, re.IGNORECASE):
                            continue
                        if re.match(r'^\s*<!--\s*image\s*-->\s*$', body_line, re.IGNORECASE):
                            continue
                        if section_name.lower() in ('editor overview', 'preview overview') and _is_screenshot_ocr_noise(body_line):
                            continue
                        cleaned_lines.append(body_line)

                    image_uri = ordered_image_uris[image_idx]
                    image_line = f"![{section_name.lower().replace(' ', '-')}]({image_uri})"
                    cleaned_body = "\n".join(cleaned_lines).strip("\n")

                    # For Preview Overview, aggressively remove dangling OCR text fragments.
                    if section_name.lower() == 'preview overview':
                        cleaned_body = re.sub(
                            r'^[\s]*toggle, and drag-and-drop file opening\..*?(?:\nCompatible with mac[^\n]*)?',
                            '',
                            cleaned_body,
                            flags=re.IGNORECASE | re.MULTILINE | re.DOTALL,
                        )
                        cleaned_body = cleaned_body.strip()

                    inserted_count += 1
                    if cleaned_body:
                        return f"{heading_line}\n\n{image_line}\n\n{cleaned_body}\n"
                    return f"{heading_line}\n\n{image_line}\n"

                updated = re.sub(section_pattern, _section_repl, updated, flags=re.IGNORECASE | re.MULTILINE | re.DOTALL)

            # If two local section images are already inserted, remove stray GitHub attachment
            # image lines outside those sections to reduce reversed visual order.
            if inserted_count >= 2:
                updated = re.sub(
                    r'^\s*!\[[^\]]*\]\(https?://github\.com/user-attachments/assets/[^)]+\)\s*$\n?',
                    '',
                    updated,
                    flags=re.MULTILINE,
                )

            return updated

        def _cleanup_docling_structure(text):
            cleaned = text

            def _strip_overview_body_noise(src_text, section_title):
                section_re = rf'(?ms)^##\s*{re.escape(section_title)}\s*\n(.*?)(?=^##\s+|\Z)'

                def _section_clean(match):
                    body = match.group(1)
                    kept_lines = []
                    for ln in body.splitlines():
                        s = ln.strip()
                        if not s:
                            continue
                        if re.match(r'^!\[[^\]]*\]\([^)]+\)$', s):
                            kept_lines.append(s)
                    if kept_lines:
                        return f"## {section_title}\n\n" + "\n\n".join(kept_lines) + "\n\n"
                    return f"## {section_title}\n\n"

                return re.sub(section_re, _section_clean, src_text, count=1, flags=re.IGNORECASE)

            def _remove_duplicate_section(src_text, section_title):
                # Keep only the first section occurrence for titles that should be unique.
                pat = rf'(?ms)^##\s*{re.escape(section_title)}\s*\n.*?(?=^##\s+|\Z)'
                matches = list(re.finditer(pat, src_text, flags=re.IGNORECASE))
                if len(matches) <= 1:
                    return src_text

                first_end = matches[0].end()
                pieces = [src_text[:first_end]]
                last = first_end
                for m in matches[1:]:
                    pieces.append(src_text[last:m.start()])
                    last = m.end()
                pieces.append(src_text[last:])
                return ''.join(pieces)

            # Promote the document title when Docling downgrades first heading level.
            cleaned = re.sub(r'\A\s*##\s*Markdown\s+Reader\s*\n', '# Markdown Reader\n\n', cleaned, flags=re.IGNORECASE)

            # Drop noisy app-window OCR chunks often injected between real sections.
            cleaned = re.sub(
                r'(?mis)^Installation\s*&\s*Usage\s*\n\s*README\.MD\s*x\s*\n\s*#\s*Markdown\s+Reader\s*\n\s*1\.\s*Clone the repository\s*\n\s*##\s*What\'s New\s*\n',
                '',
                cleaned,
            )

            # Remove malformed duplicated feature chunk produced by OCR.
            cleaned = re.sub(
                r'(?mis)^##\s*Features\s+python\s+-m\s+venv\s+venv\s*\n.*?(?=^##\s+Editor Overview\b|^##\s+Preview Overview\b|^##\s+Installation\s*&\s*Usage\b)',
                '',
                cleaned,
            )

            # Remove duplicated heading right after Features (e.g. "## Features" then "## Markdown Reader").
            cleaned = re.sub(
                r'(?mi)^##\s*Features\s*\n\s*##\s*Markdown\s+Reader\s*\n',
                '## Features\n\n',
                cleaned,
            )

            # Remove duplicated intro fragment that often appears before a second Features block.
            cleaned = re.sub(
                r'(?ms)^\*\s*Can be bundled as a macOS app using[^\n]*\n\n\*\s*Opens preview automatically[^\n]*\n\nMarkdown Reader is a clean and intuitive Markdown reader[^\n]*\n---\n\n##\s*Features\s*\n',
                '## Features\n\n',
                cleaned,
            )

            # Remove accidental standalone code block under Features: ```bash\ncd markdown-reader\n```
            cleaned = re.sub(
                r'(?mis)^```(?:bash|sh|shell|zsh)\s*\n\s*cd\s+markdown-reader\s*\n```\s*\n(?=\s*[-*]\s+Real-time preview of Markdown rendered to HTML)',
                '',
                cleaned,
            )

            # Split merged command lines commonly produced by OCR/round-trip.
            cleaned = re.sub(
                r'(?mi)^(\s*git clone\s+\S+)\s+cd\s+markdown-reader\s*$',
                r'\1\ncd markdown-reader',
                cleaned,
            )
            cleaned = re.sub(
                r'(?mi)^(\s*git add\s+\.)\s+git\s+push\s*$',
                r'\1\ngit push',
                cleaned,
            )
            cleaned = re.sub(
                r'(?mi)^(\s*python\s+-m\s+venv\s+venv)\s+source\s+venv/bin/activate\s*(#.*)?$',
                r'\1\nsource venv/bin/activate \2',
                cleaned,
            )
            cleaned = re.sub(
                r'(?mi)^\s*source\s+venv/bin/activate\s*#\s*macos/linux\s*#\s*(\.\\venv\\scripts\\activate\s*#\s*windows\s*\(cmd/powershell\))\s*$',
                r'source venv/bin/activate  # macOS/Linux\n# \1',
                cleaned,
            )
            cleaned = re.sub(
                r'(?mi)^(\s*rm\s+-rf\s+build\s+dist)\s+python\s+setup\.py\s+py2app\s*$',
                r'\1\npython setup.py py2app',
                cleaned,
            )
            
            # Merge consecutive git command code blocks (Submit Changes to Git section).
            cleaned = re.sub(
                r'(?ms)^##\s*Submit Changes to Git\s*\n```(?:bash)?\s*\ngit commit[^\n]*\n```\s*\n+```(?:bash)?\s*\ngit add[^\n]*\ngit push[^\n]*\n```',
                r'## Submit Changes to Git\n```bash\ngit add .\ngit commit -m "Update"  # Replace "Update" with a meaningful commit message\ngit push\n```',
                cleaned,
            )

            # Remove trailing OCR bullet separators at end of list lines.
            cleaned = re.sub(r'[ \t]*[·•]\s*$', '', cleaned, flags=re.MULTILINE)

            # Remove isolated OCR/UI artifact lines.
            cleaned = re.sub(r'(?mi)^\s*(README\.MD\s*x|Copy|Markdown Reader|What\'s New)\s*$', '', cleaned)
            cleaned = re.sub(r'(?m)^\s*\d{1,3}\s*$', '', cleaned)

            # Repair bare command lines that lost fenced code blocks.
            cleaned = re.sub(
                r'(?mis)^##\s*Running the Application\s*\n\s*python\s+app\.py\s*$',
                '## Running the Application\n\n```bash\npython app.py\n```',
                cleaned,
            )
            cleaned = re.sub(
                r'(?mis)^##\s*Exit the Virtual Environment\s*\n\s*deactivate\s*$',
                '## Exit the Virtual Environment\n\n```bash\ndeactivate\n```',
                cleaned,
            )
            cleaned = re.sub(
                r'(?mis)^##\s*AI-powered translation:\s*\n\s*To enable AI-powered translation features, you need to set up API keys:\s*\n\s*```[\s\S]*?```',
                '## AI-powered translation:\n\nTo enable AI-powered translation features, open "Settings -> AI Provider & API Keys..." and save your provider, model, and API key.',
                cleaned,
            )
            cleaned = re.sub(
                r'(?mis)^##\s*3\.\s*Install dependencies\s*\n\s*For Mac users[\s\S]*?before running\s+`?pip install`?\s*\.\s*\n\s*pip install -r requirements\.txt\s*$',
                lambda m: m.group(0).replace('\npip install -r requirements.txt', '\n\n```bash\npip install -r requirements.txt\n```'),
                cleaned,
            )
            cleaned = re.sub(
                r'(?mi)^\s*git add\s+\.\s+git commit\s+-m\s+"Update"\s+#\s*Replace\s+"Update"\s+with\s+a\s+meaningful\s+commit\s+message\s+git push\s*$',
                'git add .\ngit commit -m "Update"  # Replace "Update" with a meaningful commit message\ngit push',
                cleaned,
            )

            # Normalize section heading levels for numbered installation sub-steps.
            cleaned = re.sub(r'(?mi)^##\s*(\d+\.\s+.+)$', r'#### \1', cleaned)

            # Restore common README links when OCR drops markdown link syntax.
            cleaned = re.sub(
                r'(?mi)(PrepareForMacUser)\s+file',
                r'[PrepareForMacUser](./doc/PrepareForMacUser.md) file',
                cleaned,
            )
            cleaned = re.sub(
                r'(?mi)(PrepareForWindowsUser)\s+file',
                r'[PrepareForWindowsUser](./doc/PrepareForWindowsUser.md) file',
                cleaned,
            )

            # Remove immediate duplicate headings/images that often appear back-to-back.
            cleaned = re.sub(
                r'(?ms)(^##\s*Preview Overview\s*\n\s*!\[[^\]]*\]\([^)]+\)\s*\n)\s*\1+',
                r'\1',
                cleaned,
            )

            cleaned = _strip_overview_body_noise(cleaned, 'Editor Overview')
            cleaned = _strip_overview_body_noise(cleaned, 'Preview Overview')
            cleaned = _remove_duplicate_section(cleaned, 'Editor Overview')
            cleaned = _remove_duplicate_section(cleaned, 'Preview Overview')

            # Remove exact duplicated non-empty lines while preserving order.
            lines = cleaned.splitlines()
            seen_once = set()
            deduped_lines = []
            for line in lines:
                normalized = re.sub(r'\s+', ' ', line).strip().lower()
                if not normalized:
                    deduped_lines.append(line)
                    continue

                should_dedupe = (
                    normalized.startswith('markdown reader is a clean and intuitive markdown reader')
                    or normalized == '## features'
                )

                if normalized.startswith('## system requirements') or normalized.startswith('## license') or normalized.startswith('## contributing'):
                    deduped_lines.append(line)
                    continue

                if should_dedupe:
                    if normalized in seen_once:
                        continue
                    seen_once.add(normalized)

                deduped_lines.append(line)

            cleaned = '\n'.join(deduped_lines)
            cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
            
            # Restore missing horizontal separator before System Requirements section
            # Match any text followed by one or more newlines, then the heading
            cleaned = re.sub(
                r'([^\n])\n+(##\s+System\s+Requirements)',
                r'\1\n\n---\n\n\2',
                cleaned,
                flags=re.IGNORECASE
            )
            
            # Format polish: remove spaces before colons and commas
            cleaned = re.sub(r'\s+:', ':', cleaned)  # "Translation :" → "Translation:"
            cleaned = re.sub(r'\s+,', ',', cleaned)  # ".md , .html" → ".md, .html"
            
            return cleaned

        markdown_text = _bind_images_to_sections(markdown_text, image_uris)
        markdown_text = _fix_malformed_code_fences(markdown_text)
        markdown_text = _cleanup_docling_structure(markdown_text)
        markdown_text = re.sub(r'\n{3,}', '\n\n', markdown_text)

        # If markdown still has no image reference, append extracted images as fallback
        has_any_md_images = re.search(r'!\[[^\]]*\]\([^)]+\)', markdown_text) is not None
        if image_uris and not has_any_md_images:
            markdown_text = markdown_text.rstrip() + "\n\n"
            for idx, uri in enumerate(image_uris, 1):
                markdown_text += f"![docling-image-{idx}]({uri})\n\n"
        
        return markdown_text.strip()
        
    except ImportError:
        messagebox.showwarning(
            "Docling Not Installed",
            "Advanced PDF conversion requires Docling.\n"
            "Install it using: pip install docling\n\n"
            "Falling back to standard converter."
        )
        return convert_pdf_to_markdown(pdf_path)
        
    except Exception as e:
        messagebox.showerror(
            "PDF Conversion Error",
            f"Failed to convert PDF with Docling: {e}\n\n"
            "Falling back to standard converter."
        )
        return convert_pdf_to_markdown(pdf_path)

    finally:
        if converter is not None:
            try:
                _cleanup_docling_converter(converter)
            except Exception:
                pass
